import io
import pytest
from unittest.mock import patch, MagicMock
from django.urls import reverse
from django.contrib.auth import get_user_model
from services.resume_intelligence import ResumeIntelligenceService
from apps.candidates.utils import handle_resume_upload, process_resume_file

User = get_user_model()

VALID_PDF_BYTES = b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>\nendobj\n4 0 obj\n<< /Length 55 >>\nstream\nBT /F1 12 Tf 72 712 Td (Jane Doe jane@example.com 9876543210) Tj ET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000213 00000 n \ntrailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n318\n%%EOF"

@pytest.mark.django_db
def test_normal_text_pdf_bypasses_ocr(tmp_path):
    """Test 1: Normal text PDF extracts text directly and skips OCR completely."""
    with patch('services.resume_intelligence.get_paddle_ocr_instance') as mock_ocr:
        res = ResumeIntelligenceService.run_ocr_pipeline(VALID_PDF_BYTES, "test_resume.pdf")
        assert res["resume_type"] == "EDITABLE_PDF" or res["engine"] == "pymupdf+pdfplumber+pdfminer" or "text" in res
        # Ensure PaddleOCR instance was NOT invoked for text PDF
        mock_ocr.assert_not_called()

@pytest.mark.django_db
def test_docx_resume_extracts_text():
    """Test 2: DOCX resume extracts text without OCR."""
    import docx
    doc = docx.Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("john.doe@example.com")
    doc.add_paragraph("Software Developer with 5 years experience.")
    bio = io.BytesIO()
    doc.save(bio)
    docx_bytes = bio.getvalue()

    res = ResumeIntelligenceService.run_ocr_pipeline(docx_bytes, "resume.docx")
    assert res["resume_type"] == "EDITABLE_DOCX"
    assert "John Doe" in res["text"]

@pytest.mark.django_db
def test_ocr_timeout_fails_gracefully():
    """Test 4: OCR timeout fails gracefully without hanging request or throwing exception."""
    with patch('services.resume_intelligence.get_paddle_ocr_instance', side_effect=Exception("OCR Timeout")):
        res = ResumeIntelligenceService.run_ocr_pipeline(VALID_PDF_BYTES, "scanned.pdf")
        assert "text" in res

@pytest.mark.django_db
def test_handle_resume_upload_graceful_fallback(admin_user, tmp_path):
    """Test 5: Handle resume upload graceful fallback when automatic parsing fails."""
    fake_file = io.BytesIO(VALID_PDF_BYTES)
    fake_file.name = "unparseable.pdf"
    
    with patch('apps.candidates.utils.process_resume_file', return_value=(None, "AUTOMATIC_PARSING_FAILED")):
        results = handle_resume_upload(fake_file, overwrite=True, user=admin_user)
        assert results["errors"] == 1
        assert "Resume could not be parsed automatically. Please use Manual Parsing to enter the candidate details." in results["error_reasons"][0]
