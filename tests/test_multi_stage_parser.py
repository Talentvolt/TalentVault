import pytest
from unittest.mock import MagicMock, patch
from services.parser.ocr_engine import OCREngine
from services.parser.pdf_parser import PDFParser
from services.parser.docx_parser import DOCXParser
from services.parser.layout_detector import LayoutDetector
from services.parser.table_detector import TableDetector
from docx.text.paragraph import Paragraph
from docx.table import Table
from services.parser.block_builder import BlockBuilder, NormalizedBlock, RawDocument
from services.parser.llm_extractor import LLMExtractor, save_llm_parsed_data_to_db
from django.contrib.auth import get_user_model
from apps.candidates.models import CandidateProfile

# ----------------------------------------------------------------------
# 1. OCR Engine Tests
# ----------------------------------------------------------------------
@patch('services.parser.ocr_engine.Image.open')
def test_ocr_engine_tesseract_fallback(mock_image_open):
    # Set up mock Image
    mock_image = MagicMock()
    mock_image_open.return_value = mock_image
    
    ocr = OCREngine()
    # Force PaddleOCR to fail/be disabled
    ocr._paddle_ocr = False

    # Mock pytesseract.image_to_data
    mock_data = {
        'text': ['Name:', 'John', 'Doe', ''],
        'conf': ['90', '95', '95', '-1'],
        'left': [10, 50, 100, 0],
        'top': [20, 20, 20, 0],
        'width': [30, 40, 40, 0],
        'height': [15, 15, 15, 0],
        'line_num': [1, 1, 1, 2]
    }
    
    with patch('pytesseract.image_to_data', return_value=mock_data):
        res = ocr.perform_ocr(b"fake_image_bytes")
        assert res["engine"] == "tesseract"
        assert res["text"] == "Name: John Doe"
        assert res["confidence"] == pytest.approx(93.33, 0.01)
        assert len(res["words"]) == 3
        assert res["words"][0]["text"] == "Name:"
        assert res["words"][0]["bbox"] == [10.0, 20.0, 40.0, 35.0]


# ----------------------------------------------------------------------
# 2. Table Detector Tests
# ----------------------------------------------------------------------
def test_table_detector_categories():
    td = TableDetector()

    # Education table
    edu_rows = [
        ["Degree", "Institution", "Graduation Year", "GPA"],
        ["B.Tech", "Delhi University", "2022", "8.5 CGPA"]
    ]
    assert td.detect_table_category(edu_rows) == "education"

    # Experience table
    exp_rows = [
        ["Job Title", "Company Name", "Duration"],
        ["Software Engineer", "Google", "2 years"]
    ]
    assert td.detect_table_category(exp_rows) == "experience"

    # Certification table
    cert_rows = [
        ["Certification Title", "Authority", "Credential ID"],
        ["AWS Certified Solutions Architect", "Amazon", "AWS-1234"]
    ]
    assert td.detect_table_category(cert_rows) == "certification"

    # Skills table
    skills_rows = [
        ["Programming Languages", "Proficiency Level"],
        ["Python, Go, Java", "Expert"]
    ]
    assert td.detect_table_category(skills_rows) == "skills"

    # Unknown table
    other_rows = [
        ["Item", "Price"],
        ["Apples", "$5"]
    ]
    assert td.detect_table_category(other_rows) == "other"


# ----------------------------------------------------------------------
# 3. Layout Detector Tests
# ----------------------------------------------------------------------
def test_layout_detector_header_footer():
    ld = LayoutDetector()
    
    blocks = [
        {"bbox": [20, 10, 100, 20], "text": "Header Page 1", "font_size": 8.0, "is_bold": False},
        {"bbox": [20, 100, 100, 120], "text": "WORK EXPERIENCE", "font_size": 12.0, "is_bold": True},
        {"bbox": [20, 150, 200, 180], "text": "Worked as a Software Engineer at Minda Silca.", "font_size": 10.0, "is_bold": False},
        {"bbox": [20, 780, 100, 790], "text": "Footer Info - Confidential", "font_size": 8.0, "is_bold": False}
    ]
    
    page_width = 600.0
    page_height = 800.0
    
    sorted_blocks = ld.detect_layout(blocks, page_width, page_height)
    
    # 4 blocks should remain
    assert len(sorted_blocks) == 4
    
    # Check classifications
    assert sorted_blocks[0]["layout_type"] == "header"
    assert sorted_blocks[1]["layout_type"] == "section_title"
    assert sorted_blocks[2]["layout_type"] == "body"
    assert sorted_blocks[3]["layout_type"] == "footer"


def test_layout_detector_two_columns():
    ld = LayoutDetector()
    
    # Simulate two columns: Left Column (x0=50 to x1=250), Right Column (x0=300 to x1=550)
    # The right column top block has y0=120, which is higher than left column bottom block y0=250.
    # Standard y0 sorting would interleave them: Left Top -> Right Top -> Left Bottom -> Right Bottom.
    # Correct column reading order: Left Top -> Left Bottom -> Right Top -> Right Bottom.
    blocks = [
        {"bbox": [50, 100, 250, 150], "text": "Left Col Top", "is_bold": False},
        {"bbox": [50, 250, 250, 300], "text": "Left Col Bottom", "is_bold": False},
        {"bbox": [320, 120, 550, 170], "text": "Right Col Top", "is_bold": False},
        {"bbox": [320, 270, 550, 320], "text": "Right Col Bottom", "is_bold": False}
    ]
    
    sorted_blocks = ld.detect_layout(blocks, page_width=600.0, page_height=800.0)
    
    assert len(sorted_blocks) == 4
    # Columns segregated correctly in reading order
    assert sorted_blocks[0]["text"] == "Left Col Top"
    assert sorted_blocks[1]["text"] == "Left Col Bottom"
    assert sorted_blocks[2]["text"] == "Right Col Top"
    assert sorted_blocks[3]["text"] == "Right Col Bottom"


# ----------------------------------------------------------------------
# 4. DOCX Parser Tests
# ----------------------------------------------------------------------
@patch('services.parser.docx_parser.Document')
def test_docx_parser(mock_document):
    mock_doc = MagicMock()
    mock_document.return_value = mock_doc

    # Set up mock paragraph and table child elements
    mock_para = MagicMock(spec=Paragraph)
    mock_para.text = "Job Description"
    mock_para.style.name = "Heading 1"
    mock_para.runs = []

    mock_table = MagicMock(spec=Table)
    mock_row = MagicMock()
    mock_cell = MagicMock()
    mock_cell.paragraphs = [MagicMock(text="Cell Text")]
    mock_row.cells = [mock_cell]
    mock_table.rows = [mock_row]

    # Set up XML children traversal mock
    mock_para_el = MagicMock()
    mock_para_el.tag = "w:p"
    mock_table_el = MagicMock()
    mock_table_el.tag = "w:tbl"
    
    mock_doc.element.body.iterchildren.return_value = [mock_para_el, mock_table_el]

    # Mock python-docx wrapper classes
    with patch('services.parser.docx_parser.Paragraph', return_value=mock_para), \
         patch('services.parser.docx_parser.Table', return_value=mock_table):
        
        parser = DOCXParser()
        res = parser.parse(b"fake_docx_bytes")
        
        assert len(res["blocks"]) == 2
        assert res["blocks"][0]["type"] == "text"
        assert res["blocks"][0]["text"] == "Job Description"
        assert res["blocks"][0]["style"] == "Heading 1"
        
        assert res["blocks"][1]["type"] == "table"
        assert res["blocks"][1]["text"] == "Cell Text"
        
        assert len(res["tables"]) == 1
        assert res["tables"][0]["rows"] == [["Cell Text"]]


# ----------------------------------------------------------------------
# 5. PDF Parser Tests
# ----------------------------------------------------------------------
@patch('services.parser.pdf_parser.fitz.open')
def test_pdf_parser_digital(mock_fitz_open):
    mock_doc = MagicMock()
    mock_fitz_open.return_value = mock_doc
    
    # 1 Page document
    mock_page = MagicMock()
    mock_page.rect.width = 600
    mock_page.rect.height = 800
    mock_page.find_tables.return_value = []
    mock_page.get_images.return_value = []
    
    # Mock text layout dictionary
    mock_page.get_text.return_value = {
        "blocks": [
            {
                "type": 0,
                "lines": [
                    {
                        "bbox": [50, 100, 200, 120],
                        "spans": [
                            {
                                "text": "Digital Resume Text",
                                "flags": 16,  # bold
                                "font": "Arial-Bold",
                                "size": 12.0
                            }
                        ]
                    }
                ]
            }
        ]
    }
    mock_doc.__iter__.return_value = [mock_page]
    
    parser = PDFParser()
    res = parser.parse(b"fake_pdf_bytes")
    
    assert len(res["pages"]) == 1
    page = res["pages"][0]
    assert page["is_scanned"] is False
    assert len(page["blocks"]) == 1
    assert page["blocks"][0]["text"] == "Digital Resume Text"
    assert page["blocks"][0]["is_bold"] is True
    assert page["blocks"][0]["font_size"] == 12.0


# ----------------------------------------------------------------------
# 6. Block Builder Tests
# ----------------------------------------------------------------------
def test_block_builder_pdf():
    # Setup parser mocks
    mock_pdf_parser = MagicMock()
    mock_pdf_parser.parse.return_value = {
        "pages": [
            {
                "page_index": 0,
                "width": 600.0,
                "height": 800.0,
                "blocks": [
                    {"type": "text", "bbox": [50, 100, 200, 115], "text": "Resume Summary", "font_size": 12.0, "is_bold": True}
                ],
                "tables": [
                    {"bbox": [50, 200, 550, 300], "rows": [["Company", "Role"], ["Acme Corp", "Manager"]]}
                ],
                "images": [],
                "is_scanned": False
            }
        ],
        "ocr_confidence": None
    }

    mock_table_detector = MagicMock()
    mock_table_detector.detect_table_category.return_value = "experience"

    builder = BlockBuilder(
        pdf_parser=mock_pdf_parser,
        table_detector=mock_table_detector
    )
    
    raw_doc = builder.build_raw_document(b"fake_bytes", "resume.pdf")
    
    # 2 normalized blocks (1 text summary, 1 table block formatted as text)
    assert len(raw_doc.blocks) == 2
    assert raw_doc.blocks[0].text == "Resume Summary"
    assert raw_doc.blocks[0].font_weight == "bold"
    assert raw_doc.blocks[0].type == "section_title"  # classified as title by LayoutDetector default keywords
    
    assert raw_doc.blocks[1].type == "table"
    assert "Company | Role" in raw_doc.blocks[1].text
    
    assert len(raw_doc.tables) == 1
    assert raw_doc.tables[0]["category"] == "experience"


# ----------------------------------------------------------------------
# 7. LLM Extractor & Validation Tests
# ----------------------------------------------------------------------

def test_llm_extractor_validation_engine():
    extractor = LLMExtractor()
    original_text = "John Doe is a Senior Developer at Acme Corp. Contact: john.doe@email.com"

    # Mock LLM output dictionary
    mock_extracted_data = {
        "candidate_name": {
            "value": "John Doe",
            "source_text": "John Doe is a",
            "page_number": 0,
            "confidence": 95.0
        },
        "email": {
            "value": "john.doe@email.com",
            "source_text": "Contact: john.doe@email.com",
            "page_number": 0,
            "confidence": 99.0
        },
        "phone": {
            "value": "999-888-7777",  # Hallucinated! Not in original_text
            "source_text": "Contact: 999-888-7777",
            "page_number": 0,
            "confidence": 80.0
        },
        "linkedin": {
            "value": None,
            "source_text": None,
            "page_number": 0,
            "confidence": 100.0
        },
        "github": {
            "value": None,
            "source_text": None,
            "page_number": 0,
            "confidence": 100.0
        },
        "portfolio": {
            "value": None,
            "source_text": None,
            "page_number": 0,
            "confidence": 100.0
        },
        "address": {
            "value": None,
            "source_text": None,
            "page_number": 0,
            "confidence": 100.0
        },
        "city": {
            "value": None,
            "source_text": None,
            "page_number": 0,
            "confidence": 100.0
        },
        "state": {
            "value": None,
            "source_text": None,
            "page_number": 0,
            "confidence": 100.0
        },
        "country": {
            "value": None,
            "source_text": None,
            "page_number": 0,
            "confidence": 100.0
        },
        "current_designation": {
            "value": "Senior Developer",
            "source_text": "Senior Developer at",
            "page_number": 0,
            "confidence": 95.0
        },
        "current_company": {
            "value": "Acme Corp",
            "source_text": "at Acme Corp.",
            "page_number": 0,
            "confidence": 95.0
        },
        "professional_summary": {
            "value": None,
            "source_text": None,
            "page_number": 0,
            "confidence": 100.0
        },
        "work_experience": {
            "value": [
                {
                    "company": {"value": "Acme Corp", "source_text": "Acme Corp.", "page_number": 0, "confidence": 95.0},
                    "designation": {"value": "Senior Developer", "source_text": "Senior Developer", "page_number": 0, "confidence": 95.0},
                    "location": {"value": "New York", "source_text": "New York", "page_number": 0, "confidence": 80.0}, # Hallucinated!
                    "employment_type": {"value": "Full-time", "source_text": "Full-time", "page_number": 0, "confidence": 90.0}, # Hallucinated!
                    "start_date": {"value": "2020", "source_text": "2020", "page_number": 0, "confidence": 90.0}, # Hallucinated!
                    "end_date": {"value": "Present", "source_text": "Present", "page_number": 0, "confidence": 90.0}, # Hallucinated!
                    "description": {"value": "Worked on python.", "source_text": "Worked on python.", "page_number": 0, "confidence": 90.0} # Hallucinated!
                }
            ],
            "source_text": "John Doe is...",
            "page_number": 0,
            "confidence": 95.0
        },
        "education": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "projects": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "technical_skills": {"value": ["Python", "Java"], "source_text": "Python, Java", "page_number": 0, "confidence": 90.0}, # Hallucinated! (Python, Java not in doc)
        "soft_skills": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "languages": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "certifications": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "awards": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "achievements": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "training": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "interests": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "strengths": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "references": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0}
    }

    validated = extractor.validate_document(mock_extracted_data, original_text)

    # Candidate name, email, designation, and company should be preserved (they exist in original_text)
    assert validated["candidate_name"]["value"] == "John Doe"
    assert validated["email"]["value"] == "john.doe@email.com"
    assert validated["current_designation"]["value"] == "Senior Developer"
    assert validated["current_company"]["value"] == "Acme Corp"

    # Phone is hallucinated, so it should be rejected (set to None)
    assert validated["phone"]["value"] is None
    assert validated["phone"]["source_text"] is None
    assert validated["phone"]["confidence"] == 0.0

    # Skills lists are hallucinated and should be cleaned
    assert validated["technical_skills"]["value"] == []

    # Inside work experience, company and designation are valid, but location/start_date/description are hallucinated
    exp_item = validated["work_experience"]["value"][0]
    assert exp_item["company"]["value"] == "Acme Corp"
    assert exp_item["designation"]["value"] == "Senior Developer"
    assert exp_item["location"]["value"] is None
    assert exp_item["start_date"]["value"] is None

    # Verify overall and section confidences are calculated
    assert "overall_confidence" in validated
    assert "section_confidence" in validated
    assert validated["overall_confidence"] > 0.0


@pytest.mark.django_db
def test_save_llm_parsed_data_to_db():
    # Setup test user and profile
    User = get_user_model()
    user = User.objects.create_user(email="candidate_test@example.com", password="password123")
    profile = CandidateProfile.objects.create(
        user=user,
        full_name="Original Name",
        location="Original Location",
        total_experience=1.0
    )

    mock_validated_data = {
        "candidate_name": {"value": "Ramanjeet Maurya", "source_text": "Ramanjeet Maurya", "page_number": 0, "confidence": 100.0},
        "email": {"value": "mauryaraman13@gmail.com", "source_text": "mauryaraman13@gmail.com", "page_number": 0, "confidence": 100.0},
        "phone": {"value": "9953699195", "source_text": "9953699195", "page_number": 0, "confidence": 100.0},
        "linkedin": {"value": "linkedin.com/in/ramanjeet", "source_text": "linkedin.com/in/ramanjeet", "page_number": 0, "confidence": 100.0},
        "portfolio": {"value": None, "source_text": None, "page_number": 0, "confidence": 100.0},
        "address": {"value": "Delhi, India", "source_text": "Delhi, India", "page_number": 0, "confidence": 100.0},
        "city": {"value": "Delhi", "source_text": "Delhi", "page_number": 0, "confidence": 100.0},
        "state": {"value": None, "source_text": None, "page_number": 0, "confidence": 100.0},
        "country": {"value": "India", "source_text": "India", "page_number": 0, "confidence": 100.0},
        "current_designation": {"value": "Operations Manager", "source_text": "Operations Manager", "page_number": 0, "confidence": 100.0},
        "current_company": {"value": "Amazon", "source_text": "Amazon", "page_number": 0, "confidence": 100.0},
        "professional_summary": {"value": "Experienced manager.", "source_text": "Experienced manager.", "page_number": 0, "confidence": 100.0},
        "work_experience": {
            "value": [
                {
                    "company": {"value": "Amazon", "source_text": "Amazon", "page_number": 0, "confidence": 100.0},
                    "designation": {"value": "Operations Manager", "source_text": "Operations Manager", "page_number": 0, "confidence": 100.0},
                    "location": {"value": "Delhi", "source_text": "Delhi", "page_number": 0, "confidence": 100.0},
                    "employment_type": {"value": "Full-time", "source_text": "Full-time", "page_number": 0, "confidence": 100.0},
                    "start_date": {"value": "2020-03-01", "source_text": "March 2020", "page_number": 0, "confidence": 100.0},
                    "end_date": {"value": "Present", "source_text": "Present", "page_number": 0, "confidence": 100.0},
                    "description": {"value": "- Managed site.", "source_text": "- Managed site.", "page_number": 0, "confidence": 100.0}
                }
            ],
            "source_text": "Work Experience section",
            "page_number": 0,
            "confidence": 100.0
        },
        "education": {
            "value": [
                {
                    "degree": {"value": "MBA", "source_text": "MBA", "page_number": 0, "confidence": 100.0},
                    "branch": {"value": "Operations", "source_text": "Operations", "page_number": 0, "confidence": 100.0},
                    "college": {"value": "Delhi University", "source_text": "Delhi University", "page_number": 0, "confidence": 100.0},
                    "board": {"value": None, "source_text": None, "page_number": 0, "confidence": 100.0},
                    "university": {"value": "Delhi University", "source_text": "Delhi University", "page_number": 0, "confidence": 100.0},
                    "start_year": {"value": "2013", "source_text": "2013", "page_number": 0, "confidence": 100.0},
                    "end_year": {"value": "2015", "source_text": "2015", "page_number": 0, "confidence": 100.0},
                    "cgpa": {"value": "8.5", "source_text": "8.5", "page_number": 0, "confidence": 100.0},
                    "percentage": {"value": None, "source_text": None, "page_number": 0, "confidence": 100.0},
                    "grade": {"value": "A", "source_text": "A", "page_number": 0, "confidence": 100.0}
                }
            ],
            "source_text": "Education section",
            "page_number": 0,
            "confidence": 100.0
        },
        "projects": {
            "value": [
                {
                    "title": {"value": "Tool Improvement", "source_text": "Tool Improvement", "page_number": 0, "confidence": 100.0},
                    "description": {"value": "Improved tools.", "source_text": "Improved tools.", "page_number": 0, "confidence": 100.0},
                    "technologies": {"value": "Python", "source_text": "Python", "page_number": 0, "confidence": 100.0},
                    "duration": {"value": "10 Months", "source_text": "10 Months", "page_number": 0, "confidence": 100.0}
                }
            ],
            "source_text": "Projects section",
            "page_number": 0,
            "confidence": 100.0
        },
        "technical_skills": {"value": ["Python", "Java"], "source_text": "Python, Java", "page_number": 0, "confidence": 100.0},
        "soft_skills": {"value": ["Leadership"], "source_text": "Leadership", "page_number": 0, "confidence": 100.0},
        "languages": {"value": ["English"], "source_text": "English", "page_number": 0, "confidence": 100.0},
        "certifications": {
            "value": [
                {
                    "name": {"value": "PMP Certified", "source_text": "PMP Certified", "page_number": 0, "confidence": 100.0},
                    "issuing_organization": {"value": "PMI", "source_text": "PMI", "page_number": 0, "confidence": 100.0},
                    "issue_date": {"value": "2020-05-01", "source_text": "May 2020", "page_number": 0, "confidence": 100.0}
                }
            ],
            "source_text": "Certifications section",
            "page_number": 0,
            "confidence": 100.0
        },
        "awards": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "achievements": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "training": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "interests": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "strengths": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0},
        "references": {"value": [], "source_text": None, "page_number": 0, "confidence": 100.0}
    }

    save_llm_parsed_data_to_db(profile, mock_validated_data)

    # Refresh from database
    profile.refresh_from_db()

    # Check updated fields
    assert profile.full_name == "Ramanjeet Maurya"
    assert profile.summary == "Experienced manager."
    assert profile.location == "Delhi, India"
    assert profile.linkedin_url == "linkedin.com/in/ramanjeet"
    assert profile.current_company == "Amazon"
    assert profile.current_designation == "Operations Manager"

    # Check relational tables
    assert profile.skills.count() == 3
    assert list(profile.skills.values_list("skill_name", flat=True)) == ["Python", "Java", "Leadership"]

    assert profile.experiences.count() == 1
    exp = profile.experiences.first()
    assert exp.company_name == "Amazon"
    assert exp.designation == "Operations Manager"
    assert exp.start_date.isoformat() == "2020-03-01"

    assert profile.educations.count() == 1
    edu = profile.educations.first()
    assert edu.institution == "Delhi University"
    assert edu.degree == "MBA"
    assert edu.percentage_or_cgpa == "8.5"

    assert profile.projects.count() == 1
    proj = profile.projects.first()
    assert proj.title == "Tool Improvement"

    assert profile.certifications.count() == 1
    cert = profile.certifications.first()
    assert "PMP Certified" in cert.name
    assert cert.issuing_organization == "PMI"

