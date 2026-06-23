import os
import sys
import django
import zipfile
import io
from docx import Document
from django.core.files.uploadedfile import SimpleUploadedFile
from dotenv import load_dotenv

# Add project root to sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Load env variables from .env
load_dotenv(os.path.join(os.path.dirname(__file__), '..', '.env'))

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.base')
django.setup()

from apps.candidates.utils import handle_resume_upload
from apps.accounts.models import User

# Clean up test users
User.objects.filter(email__in=["mock_docx_candidate@example.com", "zip_candidate@example.com"]).delete()

# 1. Create mock DOCX file bytes
doc = Document()
doc.add_heading('Mock Docx Candidate', 0)
doc.add_paragraph('Email: mock_docx_candidate@example.com')
doc.add_paragraph('Phone: 9999999999')
doc.add_paragraph('Work Experience: Software Developer at TechCorp')
docx_io = io.BytesIO()
doc.save(docx_io)
docx_bytes = docx_io.getvalue()

# 2. Test DOCX Upload directly
print("--- Testing DOCX file upload ---")
uploaded_docx = SimpleUploadedFile("mock_resume.docx", docx_bytes, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
results_docx = handle_resume_upload(uploaded_docx, overwrite=True)
print("DOCX Upload Results:", results_docx)

# 3. Create mock ZIP containing DOCX
print("\n--- Testing ZIP containing resume ---")
zip_io = io.BytesIO()
with zipfile.ZipFile(zip_io, 'w') as z:
    z.writestr("zip_resume.docx", docx_bytes)
zip_bytes = zip_io.getvalue()

uploaded_zip = SimpleUploadedFile("resumes.zip", zip_bytes, content_type="application/zip")
results_zip = handle_resume_upload(uploaded_zip, overwrite=True)
print("ZIP Upload Results:", results_zip)

# Clean up
User.objects.filter(email__in=["mock_docx_candidate@example.com", "zip_candidate@example.com"]).delete()
