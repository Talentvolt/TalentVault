import os
import sys
import docx

docx_path = os.path.join(os.path.dirname(__file__), '..', 'media', 'resumes', 'resume_.docx')

doc = docx.Document(docx_path)
print(f"Number of paragraphs: {len(doc.paragraphs)}")
for i, p in enumerate(doc.paragraphs[:30]):
    if p.text.strip():
        print(f"{i+1}: {p.text}")
