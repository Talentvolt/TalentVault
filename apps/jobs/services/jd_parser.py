import re
import io
import os
import logging
from typing import Dict, Any
import pdfplumber
import docx

logger = logging.getLogger(__name__)


class JobDescriptionParserService:
    @staticmethod
    def extract_text(file_bytes: bytes, filename: str) -> str:
        """
        Extract text content from uploaded PDF, DOC, or DOCX files.
        """
        ext = os.path.splitext(filename)[1].lower().replace('.', '')
        text = ""

        if ext == 'pdf':
            try:
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                logger.warning(f"[JD PARSER] PDF extraction error: {e}")

        elif ext == 'docx':
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    if para.text:
                        text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            text += row_text + "\n"
            except Exception as e:
                logger.warning(f"[JD PARSER] DOCX extraction error: {e}")

        elif ext == 'doc':
            # 1. Try docx format first (if .docx is named as .doc)
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    if para.text:
                        text += para.text + "\n"
            except Exception:
                # 2. Extract printable string sequences from binary .doc stream
                try:
                    raw_str = file_bytes.decode('latin-1', errors='ignore')
                    printable_chunks = re.findall(r'[A-Za-z0-9\s.,;:()\-+\/%\'"$@#&*]{3,}', raw_str)
                    text = "\n".join(chunk.strip() for chunk in printable_chunks if len(chunk.strip()) > 3)
                except Exception as e:
                    logger.warning(f"[JD PARSER] DOC extraction error: {e}")

        return text.strip()

    @classmethod
    def parse_jd(cls, raw_text: str) -> Dict[str, Any]:
        """
        Parse raw text into structured fields for job creation auto-fill.
        """
        if not raw_text or len(raw_text.strip()) < 10:
            return {
                "success": False,
                "error": "Unable to extract automatically. Please fill the remaining fields manually."
            }

        lines = [line.strip() for line in raw_text.split('\n') if line.strip()]

        parsed = {
            "title": "",
            "department": "",
            "job_type": "FULL_TIME",
            "location": "",
            "education": "",
            "min_experience": 0,
            "max_experience": 5,
            "min_salary": "",
            "max_salary": "",
            "required_skills": "",
            "preferred_skills": "",
            "notice_period": 30,
            "description": raw_text,
            "low_confidence_fields": []
        }

        low_conf = []

        # 1. Job Title
        title_found = False
        for line in lines[:15]:
            match = re.search(r'^(?:Job\s+Title|Title|Role|Position|Designation|Job\s+Role)\s*[:\-]\s*(.+)$', line, re.I)
            if match:
                parsed["title"] = match.group(1).strip()
                title_found = True
                break
        if not title_found:
            for line in lines[:5]:
                if len(line) < 60 and not re.search(r'^(about|overview|company|hiring|job description|page|\d+)', line, re.I):
                    parsed["title"] = line
                    title_found = True
                    break
        if not title_found or not parsed["title"]:
            parsed["title"] = ""
            low_conf.append("title")

        # 2. Department
        dept_found = False
        for line in lines[:25]:
            match = re.search(r'^(?:Department|Team|Function|Domain|Business\s+Unit)\s*[:\-]\s*(.+)$', line, re.I)
            if match:
                parsed["department"] = match.group(1).strip()
                dept_found = True
                break
        if not dept_found:
            dept_keywords = {
                "Engineering": r'\b(engineering|software|development|tech|technology|it)\b',
                "Product": r'\b(product management|product design|product)\b',
                "Data & Analytics": r'\b(data science|analytics|data engineering)\b',
                "Sales": r'\b(sales|business development|account management)\b',
                "Marketing": r'\b(marketing|growth|content)\b',
                "Human Resources": r'\b(human resources|hr|talent acquisition|recruitment)\b',
                "Finance": r'\b(finance|accounting|operations)\b',
            }
            for dept, pattern in dept_keywords.items():
                if re.search(pattern, raw_text, re.I):
                    parsed["department"] = dept
                    dept_found = True
                    break
        if not dept_found or not parsed["department"]:
            parsed["department"] = ""
            low_conf.append("department")

        # 3. Employment Type (job_type)
        type_match = re.search(r'(full[- ]time|part[- ]time|contract|freelance|hybrid|work from home|remote|wfh)', raw_text, re.I)
        if type_match:
            t = type_match.group(1).lower()
            if 'part' in t:
                parsed["job_type"] = "PART_TIME"
            elif 'contract' in t:
                parsed["job_type"] = "CONTRACT"
            elif 'freelance' in t:
                parsed["job_type"] = "FREELANCE"
            elif 'hybrid' in t:
                parsed["job_type"] = "HYBRID"
            elif 'work from home' in t or 'wfh' in t or 'remote' in t:
                parsed["job_type"] = "WORK_FROM_HOME"
            else:
                parsed["job_type"] = "FULL_TIME"
        else:
            parsed["job_type"] = "FULL_TIME"
            low_conf.append("job_type")

        # 4. Location
        loc_found = False
        for line in lines[:30]:
            match = re.search(r'^(?:Job\s+Location|Location|Workplace|City|Base\s+Location)\s*[:\-]\s*(.+)$', line, re.I)
            if match:
                parsed["location"] = match.group(1).strip()
                loc_found = True
                break
        if not loc_found:
            cities = re.findall(r'\b(Bangalore|Bengaluru|Mumbai|Pune|Hyderabad|Gurgaon|Gurugram|Noida|Delhi|Chennai|Kolkata|Ahmedabad|Remote|San Francisco|New York|London|Singapore)\b', raw_text, re.I)
            if cities:
                unique_cities = list(dict.fromkeys([c.title() for c in cities]))
                parsed["location"] = " / ".join(unique_cities[:2])
                loc_found = True
        if not loc_found or not parsed["location"]:
            parsed["location"] = ""
            low_conf.append("location")

        # 5. Education
        edu_found = False
        for line in lines:
            match = re.search(r'^(?:Education|Qualification|Degree|Educational\s+Requirement)\s*[:\-]\s*(.+)$', line, re.I)
            if match:
                parsed["education"] = match.group(1).strip()
                edu_found = True
                break
        if not edu_found:
            edu_matches = re.findall(r'\b(B\.?E\.?|B\.?Tech\.?|M\.?Tech\.?|M\.?C\.?A\.?|B\.?S\.?|M\.?S\.?|B\.?Sc\.?|M\.?Sc\.?|MBA|Bachelor|Master|Diploma)\b', raw_text, re.I)
            if edu_matches:
                unique_edu = list(dict.fromkeys([e.upper() for e in edu_matches]))
                parsed["education"] = " / ".join(unique_edu[:3])
                edu_found = True
        if not edu_found or not parsed["education"]:
            parsed["education"] = ""
            low_conf.append("education")

        # 6. Experience
        exp_found = False
        exp_match = re.search(r'(\d+)\s*(?:-|to|\+)\s*(\d+)?\s*(?:years?|yrs?|yr)\b', raw_text, re.I)
        if exp_match:
            min_e = int(exp_match.group(1))
            max_e = int(exp_match.group(2)) if exp_match.group(2) else min_e + 3
            parsed["min_experience"] = min_e
            parsed["max_experience"] = max_e
            exp_found = True
        else:
            exp_single = re.search(r'(?:min|minimum|at least)?\s*(\d+)\s*(?:years?|yrs?|yr)\s*(?:of)?\s*experience', raw_text, re.I)
            if exp_single:
                min_e = int(exp_single.group(1))
                parsed["min_experience"] = min_e
                parsed["max_experience"] = min_e + 3
                exp_found = True
        if not exp_found:
            parsed["min_experience"] = 0
            parsed["max_experience"] = 5
            low_conf.append("min_experience")
            low_conf.append("max_experience")

        # 7. Salary
        sal_found = False
        sal_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:-|to)\s*(\d+(?:\.\d+)?)\s*(?:LPA|Lacs?|Lakhs?|L|L\.P\.A\.)', raw_text, re.I)
        if sal_match:
            parsed["min_salary"] = float(sal_match.group(1))
            parsed["max_salary"] = float(sal_match.group(2))
            sal_found = True
        if not sal_found:
            parsed["min_salary"] = ""
            parsed["max_salary"] = ""
            low_conf.append("min_salary")
            low_conf.append("max_salary")

        # 8. Required & Preferred Skills
        common_tech = [
            "Python", "Java", "JavaScript", "TypeScript", "React", "Angular", "Vue", "Node.js",
            "Express", "Django", "Flask", "Spring Boot", "C++", "C#", ".NET", "PHP", "Laravel",
            "SQL", "PostgreSQL", "MySQL", "MongoDB", "Redis", "Docker", "Kubernetes", "AWS",
            "Azure", "GCP", "Linux", "DevOps", "Terraform", "Ansible", "CI/CD", "Git", "REST API",
            "GraphQL", "Microservices", "System Design", "Agile", "Scrum", "HTML", "CSS", "Tailwind"
        ]
        found_skills = []
        for skill in common_tech:
            if re.search(r'\b' + re.escape(skill) + r'\b', raw_text, re.I):
                if skill not in found_skills:
                    found_skills.append(skill)
        if found_skills:
            parsed["required_skills"] = ", ".join(found_skills[:6])
            if len(found_skills) > 6:
                parsed["preferred_skills"] = ", ".join(found_skills[6:11])
        else:
            low_conf.append("required_skills")

        # 9. Notice Period
        np_match = re.search(r'(?:Notice\s+Period|Notice)\s*[:\-]?\s*(\d+)\s*(?:days?|months?)', raw_text, re.I)
        if np_match:
            val = int(np_match.group(1))
            if 'month' in np_match.group(0).lower():
                val = val * 30
            parsed["notice_period"] = val
        else:
            if re.search(r'immediate', raw_text, re.I):
                parsed["notice_period"] = 15
            else:
                parsed["notice_period"] = 30
                low_conf.append("notice_period")

        parsed["low_confidence_fields"] = low_conf

        # Check if parsing was completely uninformative
        if len(low_conf) >= 9 and not parsed["title"] and not parsed["required_skills"]:
            return {
                "success": False,
                "error": "Unable to extract automatically. Please fill the remaining fields manually."
            }

        return {
            "success": True,
            "data": parsed
        }
