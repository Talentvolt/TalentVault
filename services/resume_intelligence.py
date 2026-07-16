import re
import os
import io
import hashlib
import numpy as np
import math
import json
import logging

logger = logging.getLogger(__name__)
from datetime import datetime
from decimal import Decimal
from django.conf import settings
from django.core.files.base import ContentFile
from apps.accounts.models import User
from apps.candidates.models import (
    CandidateProfile, CandidateSkill, Experience, Education, Project, Certification
)

import importlib.util

PADDLE_AVAILABLE = False
GLOBAL_PADDLE_OCR = None
_OCR_CACHE = {}

if importlib.util.find_spec("paddleocr") is not None:
    try:
        from paddleocr import PaddleOCR
        # Try to import paddle module as well to verify installation
        import paddle
        PADDLE_AVAILABLE = True
        GLOBAL_PADDLE_OCR = PaddleOCR(use_textline_orientation=True, lang='en')
        logger.info("Loaded global PaddleOCR instance at startup successfully.")
    except Exception as e:
        logger.error(f"Failed to load global PaddleOCR at startup: {e}")
        PADDLE_AVAILABLE = False


EASY_AVAILABLE = False
if importlib.util.find_spec("easyocr") is not None:
    try:
        import easyocr
        EASY_AVAILABLE = True
    except Exception:
        EASY_AVAILABLE = False

TESSERACT_AVAILABLE = False
if importlib.util.find_spec("pytesseract") is not None:
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        TESSERACT_AVAILABLE = True
    except Exception:
        TESSERACT_AVAILABLE = False

SPACY_AVAILABLE = False
if importlib.util.find_spec("spacy") is not None:
    try:
        import spacy
        spacy.load("en_core_web_sm")
        SPACY_AVAILABLE = True
    except Exception:
        SPACY_AVAILABLE = False


def escape_plain_text(text: str) -> str:
    if not text:
        return ""
    from xml.sax.saxutils import escape
    return escape(text)


def sanitize_html_for_reportlab(html_input) -> str:
    from xml.sax.saxutils import escape
    from bs4 import BeautifulSoup, NavigableString, Tag

    if not html_input:
        return ""

    if isinstance(html_input, str):
        if '<' not in html_input and '&' not in html_input:
            return html_input
        try:
            soup = BeautifulSoup(html_input, 'html.parser')
        except Exception:
            return escape(html_input)
    else:
        soup = html_input

    # Clean up non-content block elements
    try:
        if hasattr(soup, 'find_all'):
            for tag in soup.find_all(['style', 'script', 'head', 'title', 'meta', 'link']):
                tag.decompose()
    except Exception:
        pass

    def process_node(node) -> str:
        if isinstance(node, NavigableString):
            return escape(str(node))

        if isinstance(node, Tag):
            name = node.name.lower()

            if name in ('strong', 'b'):
                return f"<b>{process_children(node)}</b>"
            elif name in ('em', 'i'):
                return f"<i>{process_children(node)}</i>"
            elif name in ('u', 'ins'):
                return f"<u>{process_children(node)}</u>"
            elif name in ('strike', 'del', 's'):
                return f"<strike>{process_children(node)}</strike>"
            elif name == 'sub':
                return f"<sub>{process_children(node)}</sub>"
            elif name == 'sup':
                return f"<sup>{process_children(node)}</sup>"
            elif name == 'a':
                href = node.get('href', '')
                if href:
                    return f"<a href='{escape(href)}'>{process_children(node)}</a>"
                return process_children(node)
            elif name == 'font':
                attrs = []
                if node.get('color'):
                    attrs.append(f"color='{escape(node.get('color'))}'")
                if node.get('face'):
                    attrs.append(f"face='{escape(node.get('face'))}'")
                if node.get('size'):
                    attrs.append(f"size='{escape(node.get('size'))}'")
                attrs_str = " " + " ".join(attrs) if attrs else ""
                return f"<font{attrs_str}>{process_children(node)}</font>"
            elif name == 'br':
                return "<br/>"
            elif name in ('p', 'div', 'li', 'ul', 'ol', 'span'):
                return process_children(node)
            else:
                return process_children(node)

        return ""

    def process_children(tag) -> str:
        return "".join(process_node(child) for child in tag.children)

    try:
        result = process_children(soup) if hasattr(soup, 'children') else process_node(soup)
        result = result.strip()
        while result.endswith("<br/>"):
            result = result[:-5].strip()
        return result
    except Exception:
        try:
            if hasattr(soup, 'get_text'):
                return escape(soup.get_text())
            return escape(str(soup))
        except Exception:
            return ""


def extract_bullets_from_html(html_text: str) -> list[str]:
    if not html_text:
        return []

    if '<' not in html_text:
        return [line.strip() for line in html_text.split('\n') if line.strip()]

    from bs4 import BeautifulSoup
    from xml.sax.saxutils import escape

    try:
        soup = BeautifulSoup(html_text, 'html.parser')
        
        # Remove any non-content blocks
        for tag in soup.find_all(['style', 'script', 'head', 'title', 'meta', 'link']):
            tag.decompose()

        lis = soup.find_all('li')
        if lis:
            bullets = []
            for li in lis:
                cleaned = sanitize_html_for_reportlab(li)
                if cleaned.strip():
                    bullets.append(cleaned.strip())
            return bullets

        paragraphs = soup.find_all(['p', 'div'])
        if paragraphs:
            bullets = []
            for p in paragraphs:
                if not p.find(['p', 'div']):
                    cleaned = sanitize_html_for_reportlab(p)
                    if cleaned.strip():
                        bullets.append(cleaned.strip())
            if bullets:
                return bullets

        for br in soup.find_all('br'):
            br.replace_with('\n')

        lines = soup.get_text().split('\n')
        bullets = []
        for line in lines:
            if line.strip():
                bullets.append(escape(line.strip()))
        return bullets
    except Exception:
        import re
        stripped = re.sub(r'<[^>]+>', '\n', html_text)
        return [escape(line.strip()) for line in stripped.split('\n') if line.strip()]


class ResumeIntelligenceService:
    """
    Production-grade Resume Intelligence Engine supporting multiple file formats,
    multi-engine OCR with fallback logic, Layout/NLP parsing, AI Assist, 
    and duplicate candidate similarity verification.
    """

    @staticmethod
    def parse_experience_description_to_html(desc_text: str) -> str:
        """
        Converts plain-text experience description to HTML, preserving the original
        structure exactly: bullet lines become <li> items in order, plain paragraph
        lines become <p> blocks. Never reorders lines or keyword-categorises them.
        """
        if not desc_text or not desc_text.strip():
            return ""

        # Already HTML — return unchanged
        if any(tag in desc_text for tag in ("<ul", "<li", "<p>", "<p ", "<div", "<strong")):
            style_block = (
                "<style>\n"
                "  .resume-bullets {\n"
                "    list-style-type: disc !important;\n"
                "    padding-left: 20px !important;\n"
                "    margin-bottom: 8px !important;\n"
                "  }\n"
                "  .resume-bullets li {\n"
                "    margin-bottom: 4px !important;\n"
                "    display: list-item !important;\n"
                "  }\n"
                "  .resume-ordered {\n"
                "    list-style-type: decimal !important;\n"
                "    padding-left: 20px !important;\n"
                "    margin-bottom: 8px !important;\n"
                "  }\n"
                "  .resume-ordered li {\n"
                "    margin-bottom: 4px !important;\n"
                "    display: list-item !important;\n"
                "  }\n"
                "</style>"
            )
            if "<style>" not in desc_text and ("<ul" in desc_text or "<ol" in desc_text):
                return style_block + "\n" + desc_text
            return desc_text

        BULLET_CHARS = ('-', '•', '*', '+', '●', '■', '▪', '–', '→')

        style_block = (
            "<style>\n"
            "  .resume-bullets {\n"
            "    list-style-type: disc !important;\n"
            "    padding-left: 20px !important;\n"
            "    margin-bottom: 8px !important;\n"
            "  }\n"
            "  .resume-bullets li {\n"
            "    margin-bottom: 4px !important;\n"
            "    display: list-item !important;\n"
            "  }\n"
            "  .resume-ordered {\n"
            "    list-style-type: decimal !important;\n"
            "    padding-left: 20px !important;\n"
            "    margin-bottom: 8px !important;\n"
            "  }\n"
            "  .resume-ordered li {\n"
            "    margin-bottom: 4px !important;\n"
            "    display: list-item !important;\n"
            "  }\n"
            "</style>"
        )

        lines = desc_text.split('\n')
        html_parts = []
        list_type = None  # Can be None, 'ul', or 'ol'

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                # Close any open list on blank line
                if list_type == 'ul':
                    html_parts.append('</ul>')
                elif list_type == 'ol':
                    html_parts.append('</ol>')
                list_type = None
                
                html_parts.append('<p class="mb-1">&nbsp;</p>')
                continue

            is_bullet = line[0] in BULLET_CHARS if line else False
            is_numbered = bool(re.match(r'^\d+[\.\)]\s+', line)) if line else False

            if is_bullet:
                # Strip the leading bullet character(s), preserve the rest verbatim
                content = re.sub(r'^[\-•*+●■▪–→]+\s*', '', line, count=1).strip()
                if list_type == 'ol':
                    html_parts.append('</ol>')
                    list_type = None
                if not list_type:
                    html_parts.append("<ul class='resume-bullets'>")
                    list_type = 'ul'
                html_parts.append(f'  <li>{content}</li>')
            elif is_numbered:
                # Strip the leading numbering pattern (e.g. 1. or 2)), preserve the rest
                content = re.sub(r'^\d+[\.\)]\s*', '', line, count=1).strip()
                if list_type == 'ul':
                    html_parts.append('</ul>')
                    list_type = None
                if not list_type:
                    html_parts.append("<ol class='resume-ordered'>")
                    list_type = 'ol'
                html_parts.append(f'  <li>{content}</li>')
            else:
                # Plain paragraph line
                if list_type == 'ul':
                    html_parts.append('</ul>')
                elif list_type == 'ol':
                    html_parts.append('</ol>')
                list_type = None
                
                html_parts.append(f'<p class="mb-1">{line}</p>')

        # Close any trailing open list
        if list_type == 'ul':
            html_parts.append('</ul>')
        elif list_type == 'ol':
            html_parts.append('</ol>')

        joined_html = '\n'.join(html_parts)
        if "class='resume-bullets'" in joined_html or "class='resume-ordered'" in joined_html:
            return style_block + "\n" + joined_html
        return joined_html

    @staticmethod
    def is_valid_name(name: str) -> bool:
        if not name or not isinstance(name, str):
            return False
        name_clean = " ".join(name.strip().split())
        if not name_clean:
            return False
            
        # Reject digits/phone patterns
        if name_clean.isdigit():
            return False
        if re.match(r'^\+?\d[\d\s-]{8,}$', name_clean):
            return False
        if '@' in name_clean:
            return False
        if name_clean.lower().startswith('http'):
            return False
        if 'linkedin' in name_clean.lower() or 'github' in name_clean.lower():
            return False
            
        digits_only = re.sub(r'[^\d+]', '', name_clean)
        if len(digits_only) >= 8 and digits_only.replace('+', '').isdigit():
            return False
            
        if re.search(r'[\w\.-]+@[\w\.-]+\.\w+', name_clean):
            return False
        if re.search(r'(https?://\S+|www\.\S+)', name_clean, re.I):
            return False
            
        if not any(char.isalpha() for char in name_clean):
            return False
            
        # Normalize and reject section titles
        norm = re.sub(r'[^a-z\s]', '', name_clean.lower()).strip()
        norm = " ".join(norm.split())
        
        SECTION_TITLES = {
            "objective", "summary", "professional summary", "profile", "education",
            "experience", "work experience", "projects", "technical skills", "skills",
            "certifications", "achievements", "awards", "languages", "personal details",
            "interests", "hobbies", "extracurricular activities", "volunteer work",
            "declaration", "references", "career objective", "academic qualification"
        }
        if norm in SECTION_TITLES:
            return False
            
        common_headings = {
            'curriculum vitae', 'curriculum', 'vitae', 'resume', 'cv', 'biodata', 'page', 'email', 'phone', 'contact', 'mobile'
        }
        if norm in common_headings:
            return False
            
        # Reject standalone blacklisted words
        words = name_clean.lower().split()
        blacklisted_words = {
            'manager', 'developer', 'executive', 'engineer', 'lead', 'associate', 'specialist', 'director', 
            'analyst', 'consultant', 'officer', 'administrator', 'coordinator', 'technician', 'representative', 
            'intern', 'programmer', 'architect', 'head', 'founder', 'co-founder', 'ceo', 'cto', 'supervisor',
            'leader', 'operator', 'agent', 'strategist', 'advisor', 'expert', 'auditor', 'salesperson',
            'ltd', 'limited', 'pvt', 'private', 'llp', 'llc', 'inc', 'company', 'corporation', 'technologies',
            'solutions', 'industries', 'group', 'corp', 'hospital', 'university', 'college', 'institute',
            'school', 'bank', 'unknown', 'hometown', 'residence', 'nationality', 'gender', 'about', 'hr',
            'recruiter', 'team', 'page', 'phone', 'email', 'address', 'contact', 'mobile', 'cv', 'resume',
            'biodata', 'curriculum', 'vitae'
        }
        if any(w in blacklisted_words for w in words):
            return False
            
        # Must not be a single long run-on word without spaces (e.g. CURRICULUMVITAE)
        if ' ' not in name_clean and len(name_clean) > 12:
            return False
            
        # Must have between 2 and 5 words
        if not (2 <= len(words) <= 5):
            return False
            
        return True

    @staticmethod
    def calculate_name_email_similarity(name: str, email: str, linkedin: str = "") -> float:
        # Extract username from email
        username = email.split('@')[0].lower() if email else ""
        # Extract username from LinkedIn URL
        li_user = ""
        if linkedin:
            parts = linkedin.strip('/').split('/')
            if parts:
                li_user = parts[-1].lower()
                
        name_words = re.sub(r'[^a-zA-Z\s]', '', name).lower().split()
        if not name_words:
            return 0.0
            
        score = 0.0
        # Check email match
        if username:
            clean_username = re.sub(r'[^a-z]', ' ', username)
            user_words = clean_username.split()
            for w in name_words:
                if w in user_words or any(w in uw or uw in w for uw in user_words):
                    score += 0.5
                    
        # Check LinkedIn match
        if li_user:
            clean_li = re.sub(r'[^a-z]', ' ', li_user)
            li_words = clean_li.split()
            for w in name_words:
                if w in li_words or any(w in lw or lw in w for lw in li_words):
                    score += 0.5
                    
        return score

    @staticmethod
    def detect_heading_type(line_str: str) -> str:
        l = line_str.strip().lower()
        # Remove leading/trailing bullet points or formatting
        l = re.sub(r'^[•\-\*\+\s\.,●■#]+|[•\-\*\+\s\.,●■#]+$', '', l).strip()
        l = re.sub(r'[:\-\s]+$', '', l).strip()
        if len(l) > 60 or len(l) < 3:
            return None
        
        # Exact or close matching
        if l in ["profile summary", "professional summary", "summary", "career objective", "objective", "profile", "about me", "career profile", "executive summary"]:
            return "SUMMARY"
        if l in ["work experience", "experience", "employment history", "work history", "professional experience", "professional history", "employment", "career history"]:
            return "WORK"
        if l in ["education", "academic qualification", "academic qualifications", "academic background", "qualification", "qualifications", "education history"]:
            return "EDU"
        if l in ["skills", "technical skills", "core skills", "key skills", "technical expertise", "core competencies", "expertise", "competencies", "skills & expertise"]:
            return "SKILLS"
        if l in ["projects", "personal projects", "academic projects", "key projects", "recent projects"]:
            return "PROJECT"
        if l in ["certifications", "certification", "courses", "credentials", "licenses & certifications", "certifications & licenses"]:
            return "CERT"
        if l in ["languages", "languages known", "language profile"]:
            return "LANGUAGES"
        if l in ["achievements", "key achievements", "notable achievements", "career achievements"]:
            return "ACHIEVEMENTS"
        if l in ["training", "trainings", "industrial training", "vocational training"]:
            return "TRAINING"
        if l in ["awards", "honors & awards", "honors", "accolades"]:
            return "AWARDS"
        if l in ["publications", "research publications", "patents & publications"]:
            return "PUBLICATIONS"
        if l in ["volunteer", "volunteer work", "volunteering", "social work"]:
            return "VOLUNTEER"
        if l in ["extracurricular", "extracurricular activities", "co-curricular activities", "interests", "hobbies"]:
            return "EXTRACURRICULAR"
        if l in ["personal details", "personal profile", "personal summary", "personal info"]:
            return "PERSONAL"
            
        return None

    @staticmethod
    def _date_range_regex():
        date_token = (
            r'(?:\d{1,2}[-/]\d{2,4}|'
            r'19\d\d|20\d\d|'
            r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*[-/ ]?\d{2,4}|'
            r'present|current|today|now|ongoing)'
        )
        return re.compile(
            rf'\b({date_token})\b\s*(?:-|--|to|till|until|\u2013|\u2014)\s*\b({date_token})\b',
            re.IGNORECASE
        )

    @staticmethod
    def _strip_bullet(text_value: str) -> str:
        if not text_value:
            return ""
        text_value = text_value.replace('â€¢', ' ').replace('â—', ' ').replace('â– ', ' ')
        return re.sub(r'^[\s\-*+\u2022\u25cf\u25a0]+', '', text_value).strip()

    @staticmethod
    def _looks_like_responsibility(line_str: str, responsibility_keywords: set) -> bool:
        cleaned = ResumeIntelligenceService._strip_bullet(line_str)
        if not cleaned:
            return False
        if line_str.strip().startswith(('-', '*', '+', '\u2022', '\u25cf', '\u25a0')):
            return True
        words = re.sub(r'[^a-zA-Z\s]', ' ', cleaned).lower().split()
        first_word = words[0] if words else ""
        lower = cleaned.lower()
        responsibility_headings = [
            "key result areas", "responsibilities", "key responsibilities", "duties",
            "roles and responsibilities", "role and responsibilities", "achievements"
        ]
        return first_word in responsibility_keywords or any(h in lower for h in responsibility_headings)

    @staticmethod
    def _clean_job_header_value(value: str) -> str:
        if not value:
            return ""
        value = ResumeIntelligenceService._strip_bullet(value)
        value = re.sub(r'^(presently|currently|actively)?\s*working\s+as\s+', '', value, flags=re.I).strip()
        value = re.sub(r'^worked\s+as\s+', '', value, flags=re.I).strip()
        value = re.sub(r'^(role|designation|position|title|company|organisation|organization)\s*:\s*', '', value, flags=re.I).strip()
        value = re.sub(r'^[|/\-\s\.,:]+|[|/\-\s\.,:]+$', '', value).strip()
        return " ".join(value.split())

    @staticmethod
    def _split_inline_job_header(header_text: str, designation_keywords: set, company_keywords: set) -> tuple:
        header_text = ResumeIntelligenceService._clean_job_header_value(header_text)
        if not header_text:
            return "", "", ""

        patterns = [
            r'^(?P<designation>.+?)\s+(?:in|at|with|for|@)\s+(?P<company>.+)$',
            r'^(?P<designation>.+?)\s+\|\s+(?P<company>.+)$',
            r'^(?P<designation>.+?)\s+-\s+(?P<company>.+)$',
        ]
        for pattern in patterns:
            match = re.match(pattern, header_text, re.I)
            if match:
                designation = ResumeIntelligenceService._clean_job_header_value(match.group("designation"))
                company = ResumeIntelligenceService._clean_job_header_value(match.group("company"))
                return designation, company, ""

        parts = [ResumeIntelligenceService._clean_job_header_value(p) for p in re.split(r'\s{2,}|\t+', header_text) if p.strip()]
        if len(parts) >= 2:
            scored = []
            for idx, part in enumerate(parts):
                words = set(re.sub(r'[^a-zA-Z\s]', ' ', part).lower().split())
                scored.append((idx, part, bool(words & designation_keywords), bool(words & company_keywords)))
            designation = next((p for _, p, has_desig, _ in scored if has_desig), parts[0])
            company = next((p for _, p, _, has_company in scored if has_company), "")
            if not company:
                company = next((p for _, p, _, _ in scored if p != designation), "")
            location = next((p for _, p, _, _ in scored if p not in {designation, company}), "")
            return designation, company, location

        return header_text, "", ""

    @staticmethod
    def _parse_work_experience_lines(work_lines: list) -> list:
        date_range_regex = ResumeIntelligenceService._date_range_regex()
        designation_keywords = {
            'manager', 'developer', 'executive', 'engineer', 'lead', 'associate', 'specialist', 'director',
            'analyst', 'consultant', 'officer', 'administrator', 'coordinator', 'technician', 'representative',
            'intern', 'programmer', 'architect', 'head', 'founder', 'co-founder', 'ceo', 'cto', 'supervisor',
            'leader', 'operator', 'agent', 'strategist', 'advisor', 'expert', 'auditor', 'salesperson',
            'designer', 'accountant'
        }
        company_keywords = {
            'ltd', 'limited', 'pvt', 'private', 'llp', 'llc', 'inc', 'company', 'corporation',
            'technologies', 'technology', 'solutions', 'industries', 'group', 'corp', 'jewellers',
            'lifesciences', 'motocorp', 'power', 'boilers', 'drilling', 'consulting'
        }
        responsibility_keywords = {
            'managing', 'handling', 'ensuring', 'reporting', 'coordinating', 'developing', 'providing',
            'assisting', 'performing', 'working', 'responsible', 'led', 'managed', 'coordinated',
            'assisted', 'prepared', 'monitored', 'maintained', 'drove', 'ensured', 'strengthened',
            'oversaw', 'directed', 'optimized', 'tracked', 'governed', 'crafted', 'built', 'delivered',
            'reviewed', 'partnered', 'spearheaded', 'achieved', 'contributed', 'liaised', 'supporting',
            'giving', 'execute', 'implementation', 'involved', 'floor', 'inventory', 'maintain',
            'fostering', 'monitor', 'lead', 'manage', 'coordinate', 'prepare', 'perform', 'ensure',
            'assist', 'provide', 'developed', 'designed', 'implemented', 'created', 'conducted',
            'analyzed', 'served', 'orchestrated'
        }

        job_blocks = []
        current_block = []
        in_description = False

        for raw_line in work_lines:
            line_str = raw_line.strip()
            if not line_str:
                continue
            has_date = bool(date_range_regex.search(line_str))
            is_resp = ResumeIntelligenceService._looks_like_responsibility(line_str, responsibility_keywords)
            words = set(re.sub(r'[^a-zA-Z\s]', ' ', line_str).lower().split())
            has_header_signal = bool(words & designation_keywords) or bool(words & company_keywords)
            starts_like_header = bool(line_str) and line_str[0].isupper()
            is_short_header = len(line_str) <= 95 and starts_like_header and has_header_signal and not is_resp

            is_new_job = False
            if current_block:
                current_has_date = any(date_range_regex.search(item) for item in current_block)
                if has_date and (current_has_date or in_description):
                    is_new_job = True
                elif in_description and is_short_header:
                    is_new_job = True

            if is_new_job:
                job_blocks.append(current_block)
                current_block = [line_str]
                in_description = False
            else:
                current_block.append(line_str)

            if is_resp:
                in_description = True

        if current_block:
            job_blocks.append(current_block)

        experiences = []
        for block in job_blocks:
            header_lines = []
            desc_lines = []
            block_in_desc = False
            for line in block:
                if ResumeIntelligenceService._looks_like_responsibility(line, responsibility_keywords):
                    block_in_desc = True
                if block_in_desc:
                    desc_lines.append(line)
                else:
                    header_lines.append(line)

            if not header_lines and block:
                header_lines = [block[0]]
                desc_lines = block[1:]

            start_date_val = ""
            end_date_val = ""
            remaining_headers = []
            for header in header_lines:
                match = date_range_regex.search(header)
                header_clean = header
                if match:
                    start_date_val = ResumeIntelligenceService.normalize_date_to_string(match.group(1), is_end=False) or start_date_val
                    end_date_val = ResumeIntelligenceService.normalize_date_to_string(match.group(2), is_end=True) or end_date_val
                    header_clean = (header[:match.start()] + " " + header[match.end():]).strip()
                header_clean = ResumeIntelligenceService._clean_job_header_value(header_clean)
                if header_clean:
                    remaining_headers.append(header_clean)

            designation = ""
            company = ""
            location = ""
            for header in remaining_headers:
                h_designation, h_company, h_location = ResumeIntelligenceService._split_inline_job_header(header, designation_keywords, company_keywords)
                if h_designation and not designation:
                    designation = h_designation
                if h_company and not company:
                    company = h_company
                if h_location and not location:
                    location = h_location

            if not company:
                for header in remaining_headers:
                    words = set(re.sub(r'[^a-zA-Z\s]', ' ', header).lower().split())
                    if words & company_keywords and header != designation:
                        company = header
                        break
            if not company:
                for header in remaining_headers:
                    if header != designation:
                        company = header
                        break
            if not designation:
                designation = remaining_headers[0] if remaining_headers else ""

            desc_items = []
            for desc_line in desc_lines:
                cleaned = ResumeIntelligenceService._strip_bullet(desc_line)
                if cleaned:
                    desc_items.append(f"• {cleaned}")

            desig_words = set(re.sub(r'[^a-zA-Z\s]', ' ', designation).lower().split())
            comp_words = set(re.sub(r'[^a-zA-Z\s]', ' ', company).lower().split())
            if not start_date_val and not company:
                continue
            if not start_date_val and not (desig_words & designation_keywords or comp_words & company_keywords):
                continue

            experiences.append({
                "designation": ResumeIntelligenceService._clean_job_header_value(designation),
                "company": ResumeIntelligenceService._clean_job_header_value(company),
                "location": ResumeIntelligenceService._clean_job_header_value(location),
                "duration": ResumeIntelligenceService.get_duration_display(start_date_val, end_date_val) if start_date_val else "",
                "description": "\n".join(desc_items),
                "start_date": start_date_val,
                "end_date": end_date_val
            })

        return experiences

    @staticmethod
    def _skill_group_for(skill: str) -> str:
        lower = skill.lower()
        programming = {'python', 'java', 'javascript', 'typescript', 'c', 'c++', 'c#', 'php', 'ruby', 'go', 'golang', 'scala', 'kotlin', 'swift', 'sql'}
        frameworks = {'django', 'flask', 'fastapi', 'react', 'angular', 'vue', 'node', 'express', 'laravel', 'spring', 'bootstrap', 'jquery'}
        databases = {'mysql', 'postgresql', 'postgres', 'mongodb', 'oracle', 'sqlite', 'redis', 'mssql', 'sql server'}
        cloud = {'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'terraform'}
        tools = {'git', 'github', 'gitlab', 'jira', 'confluence', 'excel', 'power bi', 'tableau', 'figma', 'postman'}
        soft = {'leadership', 'communication', 'teamwork', 'problem solving', 'time management', 'adaptability'}
        if lower in programming:
            return "Programming Languages"
        if lower in frameworks:
            return "Frameworks"
        if lower in databases:
            return "Databases"
        if lower in cloud:
            return "Cloud"
        if lower in tools:
            return "Tools"
        if lower in soft:
            return "Soft Skills"
        return "Technical Skills"

    @staticmethod
    def normalize_date_to_string(date_str: str, default_year: int = None, is_end: bool = False) -> str:
        if not date_str or not isinstance(date_str, str):
            return None
        
        val = date_str.strip().lower()
        if val in ["present", "current", "today", "now", "ongoing"]:
            return "Present"
            
        # Clean up the string a bit
        val_clean = re.sub(r'[^\w\s\-/]', ' ', val).strip()
        
        # Months mapping
        months = {
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6, 
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
            'january': 1, 'february': 2, 'march': 3, 'april': 4, 'june': 6,
            'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
        }
        
        # Pattern: Month (word) and Year (4-digit)
        m_word_year = re.search(r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s\-/]*(\d{4})\b', val, re.I)
        if m_word_year:
            m_str = m_word_year.group(1).lower()
            y_str = m_word_year.group(2)
            month = months.get(m_str, 1)
            year = int(y_str)
            day = 28 if (is_end and month == 2) else (30 if is_end and month in [4,6,9,11] else (31 if is_end else 1))
            if is_end and month == 2:
                if (year % 4 == 0 and year % 100 != 0) or (year % 400 == 0):
                    day = 29
            return f"{year:04d}-{month:02d}-{day:02d}"
            
        # Pattern: MM/YYYY or MM-YYYY or M/YYYY or M-YYYY
        m_digits_year = re.search(r'\b(\d{1,2})[\s\-/]+(\d{4})\b', val)
        if m_digits_year:
            month = int(m_digits_year.group(1))
            year = int(m_digits_year.group(2))
            if 1 <= month <= 12:
                day = 28 if (is_end and month == 2) else (30 if is_end and month in [4,6,9,11] else (31 if is_end else 1))
                if is_end and month == 2:
                    if (year % 4 == 0 and year % 100 != 0) or (year % 400 == 0):
                        day = 29
                return f"{year:04d}-{month:02d}-{day:02d}"

        # Pattern: YYYY-MM-DD or YYYY/MM/DD
        m_full_date = re.search(r'\b(\d{4})[\s\-/]+(\d{1,2})[\s\-/]+(\d{1,2})\b', val)
        if m_full_date:
            year = int(m_full_date.group(1))
            month = int(m_full_date.group(2))
            day = int(m_full_date.group(3))
            if 1 <= month <= 12 and 1 <= day <= 31:
                return f"{year:04d}-{month:02d}-{day:02d}"

        # Pattern: DD-MM-YYYY or DD/MM/YYYY
        m_reverse_date = re.search(r'\b(\d{1,2})[\s\-/]+(\d{1,2})[\s\-/]+(\d{4})\b', val)
        if m_reverse_date:
            day = int(m_reverse_date.group(1))
            month = int(m_reverse_date.group(2))
            year = int(m_reverse_date.group(3))
            if 1 <= month <= 12 and 1 <= day <= 31:
                return f"{year:04d}-{month:02d}-{day:02d}"

        # Pattern: Just YYYY (4-digit)
        m_year_only = re.search(r'\b(19\d\d|20\d\d)\b', val)
        if m_year_only:
            year = int(m_year_only.group(1))
            month = 12 if is_end else 1
            day = 31 if is_end else 1
            return f"{year:04d}-{month:02d}-{day:02d}"

        if default_year:
            month = 12 if is_end else 1
            day = 31 if is_end else 1
            return f"{default_year:04d}-{month:02d}-{day:02d}"
            
        return None

    @staticmethod
    def calculate_experience_years_from_dates(start_str: str, end_str: str) -> float:
        if not start_str:
            return 0.0
        
        try:
            start_dt = datetime.strptime(start_str, "%Y-%m-%d").date()
        except Exception:
            return 0.0
            
        if not end_str or end_str.strip().lower() in ["present", "current", "today", "now", "ongoing"]:
            end_dt = datetime.now().date()
        else:
            try:
                end_dt = datetime.strptime(end_str, "%Y-%m-%d").date()
            except Exception:
                end_dt = datetime.now().date()
                
        if start_dt > end_dt:
            return 0.0
            
        delta_days = (end_dt - start_dt).days
        years = delta_days / 365.25
        if years < 0:
            return 0.0
        return round(years, 2)

    @staticmethod
    def get_duration_display(start_str: str, end_str: str) -> str:
        if not start_str:
            return ""
        try:
            start_dt = datetime.strptime(start_str, "%Y-%m-%d").date()
        except Exception:
            return ""
            
        if not end_str or end_str.strip().lower() in ["present", "current", "today", "now", "ongoing"]:
            end_dt = datetime.now().date()
        else:
            try:
                end_dt = datetime.strptime(end_str, "%Y-%m-%d").date()
            except Exception:
                end_dt = datetime.now().date()
                
        if start_dt > end_dt:
            return "0 months"
            
        diff_years = end_dt.year - start_dt.year
        diff_months = end_dt.month - start_dt.month
        total_months = diff_years * 12 + diff_months
        
        years = total_months // 12
        months = total_months % 12
        
        parts = []
        if years > 0:
            parts.append(f"{years} Year" + ("s" if years > 1 else ""))
        if months > 0:
            parts.append(f"{months} Month" + ("s" if months > 1 else ""))
        if not parts:
            return "0 months"
        return " ".join(parts)

    @staticmethod
    def clean_camel_case_name(name: str) -> str:
        if not name or not isinstance(name, str):
            return name
        name_clean = name.strip()
        # Insert space before any uppercase letter that follows a lowercase letter (e.g. RohanKumar -> Rohan Kumar)
        # or between consecutive uppercase letters followed by a lowercase (e.g. HTMLParser -> HTML Parser)
        splitted = re.sub(r'(?<=[a-z])(?=[A-Z])|(?<=[A-Z])(?=[A-Z][a-z])', ' ', name_clean)
        # Normalize multiple spaces
        return " ".join(splitted.split())

    @staticmethod
    def extract_candidate_name(text: str, parsed_name: str = None, email: str = "", linkedin: str = "") -> str:
        text_lower = text.lower()
        email_lower = email.lower() if email else ""
        linkedin_lower = linkedin.lower() if linkedin else ""
        
        # NOTE: No hardcoded name overrides here.
        # Name must always come from the CURRENT uploaded resume's OCR text only.
        # The scoring logic below (bold text, spaCy NER, email/LinkedIn matching, line order) is sufficient.

        # General deterministic layout-aware search logic
        SECTION_TITLES = {
            "objective", "summary", "professional summary", "profile", "education",
            "experience", "work experience", "projects", "technical skills", "skills",
            "certifications", "achievements", "awards", "languages", "personal details",
            "interests", "hobbies", "extracurricular activities", "volunteer work",
            "declaration", "references", "career objective", "academic qualification", "extracurricular"
        }
        
        def is_section_heading(line_val: str) -> bool:
            cleaned = re.sub(r'^[\s\d\.\-\*•●■#]*', '', line_val).strip()
            cleaned = re.sub(r'[:\-\s]*$', '', cleaned).strip()
            norm = cleaned.lower()
            if len(norm) > 60:
                return False
            for title in SECTION_TITLES:
                pattern = r'\b' + re.escape(title) + r'\b'
                if re.search(pattern, norm):
                    return True
            return False

        def is_valid_name_candidate(name_val: str) -> bool:
            if not name_val:
                return False
            name_clean = " ".join(name_val.strip().split())
            words = name_clean.split()
            # 2–4 words
            if not (2 <= len(words) <= 4):
                return False
            # alphabetic only (except dots/hyphens)
            for w in words:
                w_clean = re.sub(r'[\.\-]', '', w)
                if not w_clean.isalpha():
                    return False
            norm = name_clean.lower()
            
            # Reject section titles
            if norm in SECTION_TITLES or any(t in norm for t in SECTION_TITLES):
                return False
                
            # Reject common resume headings
            common_headings = {
                'curriculum vitae', 'curriculum', 'vitae', 'resume', 'cv', 'biodata', 'page', 'email', 'phone', 'contact', 'mobile'
            }
            if norm in common_headings or any(ch in norm for ch in common_headings):
                return False
                
            # Reject blacklisted words / company keywords / designations
            blacklisted_words = {
                'manager', 'developer', 'executive', 'engineer', 'lead', 'associate', 'specialist', 'director', 
                'analyst', 'consultant', 'officer', 'administrator', 'coordinator', 'technician', 'representative', 
                'intern', 'programmer', 'architect', 'head', 'founder', 'co-founder', 'ceo', 'cto', 'supervisor',
                'leader', 'operator', 'agent', 'strategist', 'advisor', 'expert', 'auditor', 'salesperson',
                'ltd', 'limited', 'pvt', 'private', 'llp', 'llc', 'inc', 'company', 'corporation', 'technologies',
                'solutions', 'industries', 'group', 'corp', 'hospital', 'university', 'college', 'institute',
                'school', 'bank', 'unknown', 'hometown', 'residence', 'nationality', 'gender', 'about', 'hr',
                'recruiter', 'team', 'page', 'phone', 'email', 'address', 'contact', 'mobile', 'cv', 'resume',
                'biodata', 'curriculum', 'vitae', 'activities'
            }
            for w in words:
                if w.lower() in blacklisted_words:
                    return False
            return True

        def matches_email(cand: str, email_str: str) -> bool:
            if not email_str or not cand:
                return False
            email_user = email_str.split('@')[0].lower()
            email_user_clean = re.sub(r'[^a-z]', '', email_user)
            cand_words = [re.sub(r'[^a-z]', '', w.lower()) for w in cand.split()]
            if not cand_words:
                return False
            match_count = 0
            for w in cand_words:
                if len(w) >= 3 and w in email_user_clean:
                    match_count += 1
            return match_count >= 1

        def matches_linkedin(cand: str, linkedin_str: str) -> bool:
            if not linkedin_str or not cand:
                return False
            li_user = linkedin_str.strip('/').split('/')[-1].lower()
            li_user_clean = re.sub(r'[^a-z]', '', li_user)
            cand_words = [re.sub(r'[^a-z]', '', w.lower()) for w in cand.split()]
            if not cand_words:
                return False
            match_count = 0
            for w in cand_words:
                if len(w) >= 2 and w in li_user_clean:
                    match_count += 1
            return match_count >= 1

        # Get first page text
        page_1 = text.split('\x0c')[0] if '\x0c' in text else text
        lines = [line.strip() for line in page_1.split('\n')]
        
        # Top 20% of first page
        if len(lines) < 15:
            search_lines = lines
        else:
            top_count = max(1, int(len(lines) * 0.20))
            search_lines = lines[:top_count]
        
        # Stop at section headings
        header_lines = []
        for idx, line in enumerate(search_lines):
            if not line:
                continue
            if is_section_heading(line):
                if idx > 0 or len(header_lines) > 0:
                    break
            header_lines.append(line)

        # Extract NER (spaCy)
        spacy_persons = []
        if SPACY_AVAILABLE:
            try:
                import spacy
                nlp = spacy.load("en_core_web_sm")
                doc = nlp("\n".join(header_lines))
                for ent in doc.ents:
                    if ent.label_ == "PERSON":
                        ent_text = " ".join(ent.text.strip().split())
                        if is_valid_name_candidate(ent_text):
                            spacy_persons.append(ent_text)
            except Exception as e:
                print(f"spaCy PERSON extraction failed: {e}")

        valid_parsed_name = None
        if parsed_name:
            p_name = " ".join(parsed_name.strip().split())
            if is_valid_name_candidate(p_name):
                valid_parsed_name = p_name

        # Gather all candidate names
        candidates_map = {}
        for idx, line in enumerate(header_lines):
            line_clean = " ".join(line.strip().split())
            if is_valid_name_candidate(line_clean):
                candidates_map[line_clean.lower()] = {
                    'name': line_clean,
                    'is_largest_bold': (valid_parsed_name is not None and line_clean.lower() == valid_parsed_name.lower()),
                    'is_spacy_person': (line_clean.lower() in [p.lower() for p in spacy_persons]),
                    'line_index': idx,
                    'in_header': True
                }

        if valid_parsed_name:
            key = valid_parsed_name.lower()
            if key not in candidates_map:
                candidates_map[key] = {
                    'name': valid_parsed_name,
                    'is_largest_bold': True,
                    'is_spacy_person': (key in [p.lower() for p in spacy_persons]),
                    'line_index': 0,
                    'in_header': False
                }

        for p in spacy_persons:
            key = p.lower()
            if key not in candidates_map:
                candidates_map[key] = {
                    'name': p,
                    'is_largest_bold': (valid_parsed_name is not None and key == valid_parsed_name.lower()),
                    'is_spacy_person': True,
                    'line_index': 1,
                    'in_header': False
                }

        # Calculate scores based on the ensemble priority
        candidates = list(candidates_map.values())
        for c in candidates:
            score = 0.0
            if c['is_largest_bold']:
                score += 10.0
            if c['in_header']:
                score += 5.0
            if matches_email(c['name'], email):
                score += 3.0
            if matches_linkedin(c['name'], linkedin):
                score += 2.0
            if c['is_spacy_person']:
                score += 1.0
            
            score += (1.0 / (c['line_index'] + 1))
            c['score'] = score

        if candidates:
            candidates.sort(key=lambda x: x['score'], reverse=True)
            return candidates[0]['name'].title()
            
        return "Unknown Candidate"

    @staticmethod
    def detect_resume_type(filename: str) -> str:
        ext = filename.split('.')[-1].lower() if '.' in filename else ''
        if ext in ['pdf']:
            return 'PDF'
        elif ext in ['docx']:
            return 'DOCX'
        elif ext in ['doc']:
            return 'DOC'
        elif ext in ['rtf']:
            return 'RTF'
        elif ext in ['txt']:
            return 'TXT'
        elif ext in ['zip']:
            return 'ZIP'
        elif ext in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
            return 'IMAGE'
        return 'UNKNOWN'

    @staticmethod
    def merge_extracted_texts(text1: str, text2: str, text3: str) -> str:
        text1 = text1 or ""
        text2 = text2 or ""
        text3 = text3 or ""
        
        base_text = text1
        if len(text2) > len(base_text):
            base_text = text2
        if len(text3) > len(base_text):
            base_text = text3
            
        other_texts = [t for t in [text1, text2, text3] if t != base_text]
        merged_lines = base_text.split('\n')
        base_normalized = {re.sub(r'\s+', '', line).lower() for line in merged_lines if line.strip()}
        
        for other in other_texts:
            if not other:
                continue
            for line in other.split('\n'):
                line_strip = line.strip()
                if not line_strip:
                    continue
                line_norm = re.sub(r'\s+', '', line_strip).lower()
                if line_norm not in base_normalized:
                    merged_lines.append(line_strip)
                    base_normalized.add(line_norm)
                    
        return "\n".join(merged_lines)

    @staticmethod
    def preprocess_image_for_ocr(img):
        import cv2
        import numpy as np
        from PIL import Image

        # 1. Orientation Correction (using Tesseract OSD if available)
        if TESSERACT_AVAILABLE:
            try:
                import pytesseract
                osd = pytesseract.image_to_osd(img)
                rotation = re.search(r'Rotate: (\d+)', osd)
                if rotation:
                    angle = int(rotation.group(1))
                    if angle == 90:
                        img = img.rotate(-90, expand=True)
                    elif angle == 180:
                        img = img.rotate(180, expand=True)
                    elif angle == 270:
                        img = img.rotate(90, expand=True)
            except Exception:
                pass

        # 2. Convert to OpenCV format (numpy array)
        img_np = np.array(img.convert('RGB'))
        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)

        # 3. Deskew
        try:
            coords = np.column_stack(np.where(gray < 255))
            if len(coords) > 0:
                angle = cv2.minAreaRect(coords)[-1]
                if angle < -45:
                    angle = -(90 + angle)
                else:
                    angle = -angle
                if abs(angle) < 15:
                    (h, w) = gray.shape[:2]
                    center = (w // 2, h // 2)
                    M = cv2.getRotationMatrix2D(center, angle, 1.0)
                    gray = cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
        except Exception:
            pass

        # 4. Resize if width is small (scale up for better OCR)
        (h, w) = gray.shape[:2]
        if w < 1500:
            scale = 1.5
            gray = cv2.resize(gray, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_CUBIC)

        # 5. Denoise
        try:
            gray = cv2.fastNlMeansDenoising(gray, h=10)
        except Exception:
            pass

        # 6. Sharpen
        try:
            kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
            gray = cv2.filter2D(gray, -1, kernel)
        except Exception:
            pass

        # 7. Adaptive Threshold
        try:
            gray = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        except Exception:
            pass

        return Image.fromarray(gray)

    @staticmethod
    def run_ocr_pipeline(file_bytes: bytes, filename: str) -> dict:
        """
        Runs OCR with fallback logic: PaddleOCR -> EasyOCR -> Tesseract.
        """
        import io
        import os
        import re
        import logging
        import numpy as np
        from PIL import Image

        logger = logging.getLogger(__name__)
        resume_type = ResumeIntelligenceService.detect_resume_type(filename)
        extracted_text = ""
        used_engine = "None (Direct Text Extraction)"
        confidence = 100.0
        largest_bold_name = None

        if resume_type == 'ZIP':
            import zipfile
            try:
                with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                    eligible_files = []
                    for name in z.namelist():
                        if name.startswith('__MACOSX/') or name.endswith('/') or name.split('/')[-1].startswith('.'):
                            continue
                        sub_ext = name.split('.')[-1].lower() if '.' in name else ''
                        if sub_ext in ['pdf', 'docx', 'doc', 'rtf', 'txt', 'png', 'jpg', 'jpeg']:
                            eligible_files.append(name)
                    
                    if eligible_files:
                        def sort_priority(n):
                            ext = n.split('.')[-1].lower()
                            if ext == 'pdf': return 0
                            if ext == 'docx': return 1
                            if ext == 'doc': return 2
                            if ext == 'rtf': return 3
                            if ext == 'txt': return 4
                            return 5
                        eligible_files.sort(key=sort_priority)
                        first_file = eligible_files[0]
                        sub_bytes = z.read(first_file)
                        return ResumeIntelligenceService.run_ocr_pipeline(sub_bytes, first_file)
            except Exception as e:
                logger.error(f"ZIP extraction/parsing failed: {e}", exc_info=True)
            return {
                "text": "Empty or unparseable ZIP file content",
                "engine": "zipfile-error",
                "confidence": 0.0,
                "resume_type": "UNKNOWN",
                "largest_bold_name": None
            }

        if resume_type == 'TXT':
            try:
                extracted_text = file_bytes.decode('utf-8', errors='ignore')
            except Exception as e:
                extracted_text = f"TXT Parse Error: {str(e)}"
            return {
                "text": extracted_text,
                "engine": "text-decode",
                "confidence": 100.0,
                "resume_type": "TEXT",
                "largest_bold_name": None
            }

        if resume_type == 'RTF':
            from striprtf.striprtf import rtf_to_text
            try:
                rtf_content = file_bytes.decode('utf-8', errors='ignore')
                extracted_text = rtf_to_text(rtf_content)
            except Exception as e:
                extracted_text = f"RTF Parse Error: {str(e)}"
            return {
                "text": extracted_text,
                "engine": "striprtf",
                "confidence": 100.0,
                "resume_type": "RTF",
                "largest_bold_name": None
            }

        if resume_type == 'DOC':
            import tempfile
            from utils.preview import convert_doc_to_pdf
            try:
                with tempfile.TemporaryDirectory() as tmpdir:
                    doc_path = os.path.join(tmpdir, "temp.doc")
                    with open(doc_path, "wb") as tmp_doc:
                        tmp_doc.write(file_bytes)
                    convert_doc_to_pdf(doc_path, tmpdir)
                    pdf_path = os.path.join(tmpdir, "temp.pdf")
                    if os.path.exists(pdf_path):
                        with open(pdf_path, "rb") as tmp_pdf:
                            pdf_bytes = tmp_pdf.read()
                        return ResumeIntelligenceService.run_ocr_pipeline(pdf_bytes, "temp.pdf")
            except Exception as e:
                logger.error(f"DOC to PDF text extraction failed: {e}", exc_info=True)
                extracted_text = f"DOC Parse Error: {str(e)}"
            return {
                "text": extracted_text,
                "engine": "soffice-pdf-fallback",
                "confidence": 50.0,
                "resume_type": "EDITABLE_DOC",
                "largest_bold_name": None
            }

        if resume_type == 'DOCX':
            import docx
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                extracted_text = "\n".join([p.text for p in doc.paragraphs])
                for p in doc.paragraphs[:15]:
                    cleaned_p = p.text.strip()
                    if not cleaned_p:
                        continue
                    is_bold = any(run.bold for run in p.runs)
                    style_name = getattr(p.style, 'name', None)
                    is_heading = style_name and (style_name.startswith("Heading") or style_name == "Title")
                    if (is_bold or is_heading) and ResumeIntelligenceService.is_valid_name(cleaned_p):
                        normalized = re.sub(r'[^a-z0-9]', '', cleaned_p.lower()).strip()
                        if normalized not in {
                            'workexperience', 'experience', 'curriculumvitae', 'resume', 'cv',
                            'biodata', 'profile', 'careerobjective', 'objective', 'summary',
                            'education', 'skills', 'projects', 'certifications', 'languages',
                            'personaldetails', 'languagesknown', 'corecompetencies', 'technicalskills',
                            'employmenthistory', 'workhistory', 'professionalexperience', 'aboutme',
                            'academicbackground', 'qualification', 'qualifications', 'educationhistory'
                        }:
                            largest_bold_name = cleaned_p
                            break
            except Exception as e:
                extracted_text = f"DOCX Parse Error: {str(e)}"
            return {
                "text": extracted_text,
                "engine": "python-docx",
                "confidence": 100.0,
                "resume_type": "EDITABLE_DOCX",
                "largest_bold_name": largest_bold_name
            }

        if resume_type == 'PDF':
            text_pymupdf = ""
            try:
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                pages_data = []
                for page_idx, page in enumerate(doc):
                    page_rect = page.rect
                    page_w = page_rect.width
                    blocks_dict = page.get_text("dict")
                    page_lines = []
                    
                    for b in blocks_dict.get("blocks", []):
                        if b.get("type") == 0:
                            for line in b.get("lines", []):
                                spans = line.get("spans", [])
                                if not spans: continue
                                line_text = " ".join("".join([s.get("text", "") for s in spans]).split())
                                if line_text:
                                    max_size = max(s.get("size", 0.0) for s in spans)
                                    is_bold = any("bold" in s.get("font", "").lower() or "black" in s.get("font", "").lower() or (s.get("flags", 0) & 16) for s in spans)
                                    x0, y0, x1, y1 = line.get("bbox", (0,0,0,0))
                                    page_lines.append({"text": line_text, "x0": x0, "y0": y0, "x1": x1, "y1": y1, "is_bold": is_bold, "font_size": max_size})
                    pages_data.append((page_w, page_lines))

                classified_pages = []
                total_left_len = 0
                total_right_len = 0
                work_heading_col = None
                work_heading_patterns = [r'^work\s+experience$', r'^experience$', r'^employment\s+history$', r'^work\s+history$', r'^professional\s+experience$', r'^professional\s+history$', r'^employment$', r'^career\s+history$', r'^experience\s+details$', r'^work\s+record$', r'^professional\s+background$', r'^experience\s+history$']
                
                for page_idx, (page_w, page_lines) in enumerate(pages_data):
                    left_lines, right_lines = [], []
                    split_candidates = np.arange(page_w * 0.35, page_w * 0.65, 5.0)
                    best_x, min_overlap = page_w / 2, float('inf')
                    for cx in split_candidates:
                        overlap_count = sum(1 for l in page_lines if l["x0"] < cx < l["x1"] and (l["x1"] - l["x0"]) < page_w * 0.5)
                        if overlap_count < min_overlap:
                            min_overlap, best_x = overlap_count, cx
                    has_columns = (len(page_lines) > 5) and (min_overlap < len(page_lines) * 0.22)
                    if has_columns:
                        for l in page_lines:
                            if ((l["x0"] + l["x1"]) / 2) <= best_x: left_lines.append(l)
                            else: right_lines.append(l)
                        for l in page_lines:
                            lt = l["text"].lower()
                            if any(re.match(p, lt) for p in work_heading_patterns):
                                work_heading_col = "LEFT" if ((l["x0"] + l["x1"]) / 2) <= best_x else "RIGHT"
                    classified_pages.append((has_columns, {"best_x": best_x, "left_boundary": page_w * 0.12}, page_lines))
                    if has_columns:
                        total_left_len += sum(len(l["text"]) for l in left_lines)
                        total_right_len += sum(len(l["text"]) for l in right_lines)

                right_column_first = (work_heading_col == "RIGHT") or (work_heading_col is None and total_right_len > total_left_len * 1.5)
                primary_stream, secondary_stream = [], []

                for page_idx, (page_w, page_lines) in enumerate(pages_data):
                    has_columns, page_info, raw_lines = classified_pages[page_idx]
                    if not page_lines: continue
                    if not has_columns:
                        lines = sorted(page_lines, key=lambda l: (l["y0"], l["x0"]))
                        formatted = [f"**{l['text']}**" if l["is_bold"] else l["text"] for l in lines]
                        primary_stream.append("\n".join(formatted))
                    else:
                        best_x, left_boundary = page_info["best_x"], page_info["left_boundary"]
                        sorted_lines = sorted(page_lines, key=lambda l: (l["y0"], l["x0"]))
                        classified, segments = [], []
                        for l in sorted_lines:
                            col = "FULL" if (l["x0"] < best_x < l["x1"] and (l["x1"] - l["x0"]) > page_w * 0.6) else ("LEFT" if ((l["x0"] + l["x1"]) / 2) <= best_x else "RIGHT")
                            classified.append((col, l))
                        
                        curr_type, curr_lines = None, []
                        for col, l in classified:
                            seg_type = "FULL" if col == "FULL" else "COLUMNS"
                            if seg_type == curr_type: curr_lines.append((col, l))
                            else:
                                if curr_lines: segments.append((curr_type, curr_lines))
                                curr_type, curr_lines = seg_type, [(col, l)]
                        if curr_lines: segments.append((curr_type, curr_lines))

                        p_parts, s_parts = [], []
                        for seg_type, seg_lines in segments:
                            if seg_type == "FULL":
                                p_parts.append("\n".join([f"**{l[1]['text']}**" if l[1]["is_bold"] else l[1]["text"] for l in seg_lines]))
                            else:
                                left_col_lines = [l[1] for l in seg_lines if l[0] == "LEFT"]
                                right_col_lines = [l[1] for l in seg_lines if l[0] == "RIGHT"]
                                def fmt(lines): return "\n".join([f"**{l['text']}**" if l["is_bold"] else l["text"] for l in lines])
                                l_txt, r_txt = fmt(left_col_lines), fmt(right_col_lines)
                                if right_column_first:
                                    if r_txt: p_parts.append(r_txt)
                                    if l_txt: s_parts.append(l_txt)
                                else:
                                    if l_txt: p_parts.append(l_txt)
                                    if r_txt: s_parts.append(r_txt)
                        if p_parts: primary_stream.append("\n".join(p_parts))
                        if s_parts: secondary_stream.append("\n".join(s_parts))
                text_pymupdf = "\n\n".join(primary_stream) + ("\n\n" + "\n\n".join(secondary_stream) if secondary_stream else "") + "\n"
            except Exception as e:
                logger.error(f"PyMuPDF direct extraction failed: {e}")

            # Optimization 1: Conditional PDF text extraction fallbacks
            if len(text_pymupdf.strip()) > 100:
                merged_text = text_pymupdf
            else:
                text_pdfplumber = ""
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text: text_pdfplumber += page_text + "\n"
                except Exception as e: logger.error(f"pdfplumber extraction failed: {e}")

                if len(text_pdfplumber.strip()) > 100:
                    merged_text = ResumeIntelligenceService.merge_extracted_texts(text_pdfplumber, "", "")
                else:
                    text_pdfminer = ""
                    try:
                        from pdfminer.high_level import extract_text as pdfminer_extract_text
                        text_pdfminer = pdfminer_extract_text(io.BytesIO(file_bytes))
                    except Exception as e: logger.error(f"pdfminer.six extraction failed: {e}")
                    merged_text = ResumeIntelligenceService.merge_extracted_texts(text_pdfplumber, text_pdfminer, "")

            if len(merged_text.strip()) > 100:
                return {"text": merged_text, "engine": "pymupdf+pdfplumber+pdfminer", "confidence": 99.0, "resume_type": "EDITABLE_PDF", "largest_bold_name": largest_bold_name}
            else:
                resume_type, used_engine = 'IMAGE', "Scanned PDF (OCR Required)"

        image_list = []
        try:
            if resume_type == 'PDF' or filename.lower().endswith('.pdf'):
                import pypdfium2
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                pdf = pypdfium2.PdfDocument(file_bytes)
                for i in range(len(pdf)):
                    fitz_page = doc[i]
                    total_text_len = len(fitz_page.get_text().strip())
                    
                    # Skip OCR if the page already has selectable text
                    if total_text_len >= 15:
                        image_list.append((i, "text", fitz_page.get_text()))
                    else:
                        # Skip OCR if there are no images on the page
                        has_images = len(fitz_page.get_images()) > 0
                        if has_images:
                            page = pdf[i]
                            bitmap = page.render(scale=2)
                            image_list.append((i, "image", bitmap.to_pil()))
                        else:
                            image_list.append((i, "blank", ""))
            else:
                image_list.append((0, "image", Image.open(io.BytesIO(file_bytes))))
        except Exception as e:
            logger.error(f"Image rendering/loading failed: {e}")
            return {"text": extracted_text or "Could not open document.", "engine": "Error / Fallback Text", "confidence": 50.0, "resume_type": "CORRUPTED", "largest_bold_name": None}

        # Local helper for image compression
        def compress_image_for_ocr(img, max_size=1500):
            try:
                w, h = img.size
                if not isinstance(w, (int, float)) or not isinstance(h, (int, float)):
                    return img
            except Exception:
                return img
            if max(w, h) > max_size:
                ratio = max_size / float(max(w, h))
                new_size = (int(w * ratio), int(h * ratio))
                img = img.resize(new_size, Image.Resampling.LANCZOS)
            return img

        # Local helper to run OCR on a single page
        def process_single_page_ocr(p_idx, img):
            # Duplicate OCR check (page-level cache)
            raw_bytes = img.tobytes()
            img_hash = hashlib.sha256(raw_bytes).hexdigest()
            if img_hash in _OCR_CACHE:
                cached = _OCR_CACHE[img_hash]
                return p_idx, cached["text"], cached["engine"], cached["confidence"]

            img_compressed = compress_image_for_ocr(img)
            img_preproc = ResumeIntelligenceService.preprocess_image_for_ocr(img_compressed)
            
            page_text, engine_success, used_engine_name, page_confidence = "", False, "", 0.0
            
            # Try PaddleOCR (up to 2 attempts)
            if PADDLE_AVAILABLE:
                for attempt in range(2):
                    try:
                        ocr = GLOBAL_PADDLE_OCR
                        if ocr is None:
                            from paddleocr import PaddleOCR
                            GLOBAL_PADDLE_OCR = PaddleOCR(use_textline_orientation=True, lang='en')
                            ocr = GLOBAL_PADDLE_OCR
                        res = ocr.ocr(np.array(img_preproc.convert('RGB')), cls=True)
                        if res and res[0]:
                            lines, confs = [line[1][0] for line in res[0]], [line[1][1] * 100 for line in res[0]]
                            page_text, page_confidence = "\n".join(lines), sum(confs) / len(confs)
                            used_engine_name, engine_success = "PaddleOCR", True
                            break
                    except Exception as e:
                        logger.error(f"PaddleOCR failed on page {p_idx} (attempt {attempt+1}): {e}")
            
            # Try EasyOCR (up to 2 attempts)
            if not engine_success and EASY_AVAILABLE:
                for attempt in range(2):
                    try:
                        import easyocr
                        reader = easyocr.Reader(['en'])
                        res = reader.readtext(np.array(img_preproc.convert('RGB')))
                        if res:
                            lines, confs = [r[1] for r in res], [r[2] * 100 for r in res]
                            page_text, page_confidence = "\n".join(lines), sum(confs) / len(confs)
                            used_engine_name, engine_success = "EasyOCR", True
                            break
                    except Exception as e:
                        logger.error(f"EasyOCR failed on page {p_idx} (attempt {attempt+1}): {e}")
            
            # Try Tesseract (up to 2 attempts)
            if not engine_success and TESSERACT_AVAILABLE:
                for attempt in range(2):
                    try:
                        import pytesseract
                        text = pytesseract.image_to_string(img_preproc)
                        if text.strip():
                            page_text, page_confidence, used_engine_name = text, 92.5, "Tesseract"
                            engine_success = True
                            break
                    except Exception as e:
                        logger.error(f"Tesseract failed on page {p_idx} (attempt {attempt+1}): {e}")

            res_dict = {
                "text": page_text,
                "engine": used_engine_name or "None",
                "confidence": page_confidence
            }
            _OCR_CACHE[img_hash] = res_dict
            return p_idx, page_text, used_engine_name, page_confidence

        # Run page preprocessing and OCR in parallel!
        import concurrent.futures
        ocr_results_by_page = {}
        ocr_tasks = [(idx, img) for idx, page_type, img in image_list if page_type == "image"]
        
        if ocr_tasks:
            with concurrent.futures.ThreadPoolExecutor(max_workers=min(4, len(ocr_tasks))) as executor:
                futures = {executor.submit(process_single_page_ocr, idx, img): idx for idx, img in ocr_tasks}
                for future in concurrent.futures.as_completed(futures):
                    idx = futures[future]
                    try:
                        p_idx, page_text, used_engine_name, page_confidence = future.result()
                        ocr_results_by_page[p_idx] = {
                            "text": page_text,
                            "engine": used_engine_name,
                            "confidence": page_confidence
                        }
                    except Exception as e:
                        logger.error(f"Parallel OCR failed for page {idx}: {e}")
                        ocr_results_by_page[idx] = {"text": "", "engine": "Failed", "confidence": 0.0}

        # Reassemble page results
        ocr_results = {"text": "", "engine": "", "confidence": 0.0}
        for idx, page_type, page_data in image_list:
            if page_type == "text":
                ocr_results["text"] += page_data + "\n\n"
                ocr_results["engine"] = "pymupdf"
                ocr_results["confidence"] = max(ocr_results["confidence"], 99.0)
            elif page_type == "image":
                page_res = ocr_results_by_page.get(idx, {"text": "", "engine": "", "confidence": 0.0})
                ocr_results["text"] += page_res["text"] + "\n\n"
                ocr_results["engine"] = page_res["engine"] or ocr_results["engine"]
                ocr_results["confidence"] = max(ocr_results["confidence"], page_res["confidence"])

        result_dict = {"text": ocr_results["text"].strip(), "engine": ocr_results["engine"] or "None", "confidence": ocr_results["confidence"] or 0.0, "resume_type": "SCANNED_IMAGE" if resume_type == 'IMAGE' else "SCANNED_PDF", "largest_bold_name": None}
        
        # Fallback if OCR text is empty
        if not result_dict["text"].strip():
            logger.warning("OCR returned empty text. Running direct extraction fallback (PyMuPDF, pdfplumber, pdfminer)...")
            fallback_text = ""
            try:
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    fallback_text += page.get_text() + "\n"
            except Exception as e:
                logger.error(f"Fallback PyMuPDF failed: {e}")
            
            if not fallback_text.strip():
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                fallback_text += page_text + "\n"
                except Exception as e:
                    logger.error(f"Fallback pdfplumber failed: {e}")
            
            if not fallback_text.strip():
                try:
                    from pdfminer.high_level import extract_text as pdfminer_extract_text
                    fallback_text = pdfminer_extract_text(io.BytesIO(file_bytes))
                except Exception as e:
                    logger.error(f"Fallback pdfminer failed: {e}")
            
            if fallback_text.strip():
                result_dict["text"] = fallback_text.strip()
                result_dict["engine"] = "ocr-failed-fallback-direct"
        if filename.split('.')[-1].lower() in ['png', 'jpg', 'jpeg', 'tiff', 'bmp'] or result_dict["resume_type"] == "SCANNED_IMAGE":
            text_lines = result_dict['text'].split('\n')
            first_20_lines = "\n".join(text_lines[:20])
            print(f"[OCR_CONFIDENCE] {result_dict['confidence']}\n[FIRST_20_LINES]:\n{first_20_lines}")
            logger.info(f"\n[IMAGE DETECTED] Processing file: {filename}\n[OCR ENGINE USED] {result_dict['engine']}\n[OCR CONFIDENCE] {result_dict['confidence']}\n[TEXT LENGTH] {len(result_dict['text'])}\n[FIRST 20 LINES OF EXTRACTED TEXT]:\n{first_20_lines}\n")
        return result_dict

    @staticmethod
    def parse_resume_nlp(text: str, parsed_name: str = None) -> dict:
        """
        Layout-aware Structured NLP Resume Parser that segments resumes by sections exactly,
        processes columns correctly, extracts entities with strict boundaries, and prevents misclassifications.
        """
        # Parse text into lines and identify styled lines
        raw_lines = text.split('\n')
        styled_lines = []
        for line in raw_lines:
            line_strip = line.strip()
            is_bold = line_strip.startswith('**') and line_strip.endswith('**')
            line_clean = line_strip.replace('**', '').strip()
            # Calculate indentation (count leading spaces in original line)
            indent = len(line) - len(line.lstrip())
            styled_lines.append({
                "text": line_clean,
                "is_bold": is_bold,
                "indent": indent
            })

        # Section segmentation
        current_section = "PERSONAL"
        sections = {
            "PERSONAL": [],
            "SUMMARY": [],
            "WORK": [],
            "EDU": [],
            "SKILLS": [],
            "PROJECT": [],
            "CERT": [],
            "LANGUAGES": [],
            "ACHIEVEMENTS": [],
            "OTHER": []
        }
        try:
            for line in styled_lines:
                text_val = line["text"]
                heading_type = ResumeIntelligenceService.detect_heading_type(text_val)
                if heading_type:
                    # Transition to the new section
                    current_section = heading_type if heading_type in sections else "OTHER"
                    continue
                sections[current_section].append(line)
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in section segmentation: {str(e)}", exc_info=True)

        # 1. Contact Details Regex Search
        email = ""
        phone = ""
        linkedin = ""
        portfolio = ""
        try:
            email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
            email = email_match.group(0) if email_match else ""

            phone_match = re.search(r'(?:\+?\d{1,3}[- ]?)?(?:\d[- ]?){9}\d', text)
            phone = re.sub(r'[\s-]', '', phone_match.group(0))[-10:] if phone_match else ""

            linkedin_match = re.search(r'(linkedin\.com/in/[\w-]+)', text, re.I)
            linkedin = linkedin_match.group(0) if linkedin_match else ""

            portfolio_match = re.search(r'((github\.com|portfolio|behance\.net)/[\w-]+)', text, re.I)
            portfolio = portfolio_match.group(0) if portfolio_match else ""
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in contact details extraction: {str(e)}", exc_info=True)

        # 2. Candidate Name Detection
        company_keywords = {
            'ltd', 'limited', 'pvt', 'private', 'llp', 'llc', 'inc', 'company', 'corporation',
            'jewellers', 'engineering', 'sciences', 'lifesciences', 'boilers', 'drilling', 'power',
            'motocorp', 'adani', 'systems', 'labs', 'buildgrid', 'technologies', 'solutions', 'industries', 'group'
        }
        institute_keywords = {
            'university', 'college', 'school', 'institute', 'academy', 'board', 'icai', 'icsi',
            'education', 'department', 'dept', 'polytechnic'
        }
        designation_keywords = {
            'manager', 'developer', 'executive', 'engineer', 'lead', 'associate', 'specialist', 'director',
            'analyst', 'consultant', 'officer', 'administrator', 'coordinator', 'technician', 'representative',
            'intern', 'programmer', 'architect', 'head', 'founder', 'co-founder', 'ceo', 'cto', 'supervisor',
            'leader', 'operator', 'agent', 'strategist', 'advisor', 'expert', 'auditor', 'salesperson', 'senior'
        }

        def clean_potential_name(val):
            val = val.replace('**', '').strip()
            val = re.sub(r'^[•\-\*\+\s\.,●■#]+|[•\-\*\+\s\.,●■#]+$', '', val).strip()
            return " ".join(val.split())

        def is_really_a_name(val):
            if not val or len(val) < 3 or len(val) > 50:
                return False
            if not ResumeIntelligenceService.is_valid_name(val):
                return False
            # Check keywords to avoid classifying companies/institutes/designations as candidate name
            words = set(re.sub(r'[^a-zA-Z\s]', ' ', val).lower().split())
            if words & company_keywords:
                return False
            if words & institute_keywords:
                return False
            # If it's a known designation or skill or section title
            if any(w in designation_keywords for w in words):
                # Unless it's a common name that happens to overlap (unlikely at the start of resume)
                return False
            return True

        name = ""
        try:
            if parsed_name and is_really_a_name(clean_potential_name(parsed_name)):
                name = clean_potential_name(parsed_name)
            
            if not name:
                # Look at first few lines of PERSONAL section
                for line in [l for l in sections["PERSONAL"] if l["text"]][:10]:
                    cleaned_line = clean_potential_name(line["text"])
                    if is_really_a_name(cleaned_line):
                        name = cleaned_line
                        break

            if not name:
                # Fallback check on first lines of the entire document
                for line in [l for l in styled_lines if l["text"]][:8]:
                    cleaned_line = clean_potential_name(line["text"])
                    if is_really_a_name(cleaned_line):
                        name = cleaned_line
                        break
            
            if not name:
                name = "Unknown Candidate"
            else:
                # Title-case the name for consistency (e.g. ROHAN KUMAR -> Rohan Kumar)
                name = name.strip().title()
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in candidate name detection: {str(e)}", exc_info=True)
            name = "Unknown Candidate"

        # 3. Location, Address, and City detection
        address = ""
        city = ""
        location = "Unknown"

        try:
            address_keywords = ["address", "residence", "hometown", "location"]
            # Search PERSONAL lines
            for line in sections["PERSONAL"]:
                text_val = line["text"]
                if not text_val:
                    continue
                if any(akw in text_val.lower() for akw in address_keywords):
                    addr_line = re.sub(r'^(address|residence|hometown|location)[:\-\s]*', '', text_val, flags=re.I).strip()
                    if addr_line and len(addr_line) > 3:
                        address = addr_line
                        break

            if not address:
                pincode_match = re.search(r'\b\d{6}\b', text)
                if pincode_match:
                    for line in styled_lines:
                        text_val = line["text"]
                        if text_val and pincode_match.group(0) in text_val:
                            address = text_val
                            break

            locations = ['Delhi', 'Mumbai', 'Bangalore', 'Hyderabad', 'Pune', 'Noida', 'Gurgaon', 'Patna', 'Lucknow', 'Begusarai', 'Samastipur', 'Chennai', 'Kolkata', 'Greater Noida']
            for loc in locations:
                if loc.lower() in text.lower():
                    city = loc
                    break

            if city:
                location = city
                if address and city.lower() not in address.lower():
                    address = f"{address}, {city}"
            elif address:
                location = address
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in location extraction: {str(e)}", exc_info=True)

        # 4. Professional Summary parsing
        summary = ""
        try:
            summary_lines = [l["text"] for l in sections["SUMMARY"]]
            if summary_lines:
                summary = "\n".join(summary_lines).strip()
            else:
                # Fallback: find first paragraph in PERSONAL section
                for line in sections["PERSONAL"]:
                    text_val = line["text"]
                    if len(text_val) > 80 and not any(kw in text_val.lower() for kw in ["phone", "email", "github", "linkedin", "contact"]):
                        summary = text_val
                        break
            
            if name and summary:
                summary = re.sub(rf'\b{re.escape(name)}\b', '', summary, flags=re.I).strip()
                summary = re.sub(r'^\s*[-–—,.:;]\s*', '', summary)
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in summary extraction: {str(e)}", exc_info=True)

        # 5. Parse Work Experience (Chronological, Bullets Only)
        experiences = []
        try:
            date_range_regex = ResumeIntelligenceService._date_range_regex()
            responsibility_keywords = {
                'managing', 'handling', 'ensuring', 'reporting', 'coordinating', 'developing', 'providing',
                'assisting', 'performing', 'working', 'responsible', 'led', 'managed', 'coordinated',
                'assisted', 'prepared', 'monitored', 'maintained', 'drove', 'ensured', 'strengthened',
                'oversaw', 'directed', 'optimized', 'tracked', 'governed', 'crafted', 'built', 'delivered',
                'reviewed', 'partnered', 'spearheaded', 'achieved', 'contributed', 'supporting', 'involved',
                'fostering', 'monitor', 'lead', 'manage', 'coordinate', 'prepare', 'perform', 'ensure',
                'assist', 'provide', 'developed', 'designed', 'implemented', 'created', 'conducted',
                'analyzed', 'served', 'orchestrated'
            }

            job_blocks = []
            current_block = []
            in_bullets = False

            for line in sections["WORK"]:
                text_val = line["text"]
                has_date = bool(date_range_regex.search(text_val))
                is_bullet = text_val.startswith(('-', '•', '*', '+', '●', '■')) or line["indent"] > 2
                
                # Check if this is a heading inside work section that we should skip entirely
                l_lower = text_val.lower()
                if any(h in l_lower for h in ["achievements", "certifications", "competencies", "key skills", "interests"]):
                    in_bullets = True # Ignore sub-sections in bullets
                    continue

                is_new_job = False
                if current_block:
                    current_has_date = any(date_range_regex.search(item["text"]) for item in current_block if item["text"])
                    if has_date and (current_has_date or in_bullets):
                        is_new_job = True
                    elif in_bullets and line["is_bold"] and not is_bullet and text_val:
                        # Bold line after descriptions starts a new job
                        is_new_job = True

                if is_new_job:
                    job_blocks.append(current_block)
                    current_block = [line]
                    in_bullets = False
                else:
                    current_block.append(line)

                if is_bullet:
                    in_bullets = True

            if current_block:
                job_blocks.append(current_block)

            # Sub-headings inside work sections that should trigger description mode
            work_sub_headings = {
                'key result areas', 'key result area', 'kra', 'kras',
                'roles and responsibilities', 'roles & responsibilities', 'role and responsibilities',
                'responsibilities', 'responsibility', 'duties', 'key responsibilities',
                'achievements', 'key achievements', 'notable achievements', 'accomplishments'
            }

            for block in job_blocks:
                header_lines = []
                desc_lines = []
                block_in_desc = False
                
                for line in block:
                    text_val = line["text"]
                    is_bullet = text_val.startswith(('-', '•', '*', '+', '●', '■')) or line["indent"] > 2
                    text_lower_clean = re.sub(r'[:\-\s]+$', '', text_val.strip().lower())
                    is_sub_heading = text_lower_clean in work_sub_headings
                    is_resp = is_bullet or is_sub_heading or any(w in responsibility_keywords for w in re.sub(r'[^a-zA-Z\s]', ' ', text_val).lower().split())
                    
                    if is_resp:
                        block_in_desc = True
                    
                    if block_in_desc:
                        desc_lines.append(text_val)
                    else:
                        header_lines.append(text_val)

                if not header_lines and block:
                    header_lines = [block[0]["text"]]
                    desc_lines = [b["text"] for b in block[1:]]

                # Parse dates
                start_date_val = ""
                end_date_val = ""
                remaining_headers = []
                for header in header_lines:
                    if not header:
                        continue
                    match = date_range_regex.search(header)
                    header_clean = header
                    if match:
                        start_date_val = ResumeIntelligenceService.normalize_date_to_string(match.group(1), is_end=False) or start_date_val
                        end_date_val = ResumeIntelligenceService.normalize_date_to_string(match.group(2), is_end=True) or end_date_val
                        header_clean = (header[:match.start()] + " " + header[match.end():]).strip()
                    header_clean = ResumeIntelligenceService._clean_job_header_value(header_clean)
                    if header_clean:
                        remaining_headers.append(header_clean)

                # Assign company, designation, location
                designation = ""
                company = ""
                location_val = ""

                for header in remaining_headers:
                    h_designation, h_company, h_location = ResumeIntelligenceService._split_inline_job_header(header, designation_keywords, company_keywords)
                    if h_designation and not designation:
                        designation = h_designation
                    if h_company and not company:
                        company = h_company

                if not company and designation:
                    if desc_lines:
                        first_desc = desc_lines[0].strip()
                        first_desc_clean = re.sub(r'^[•\-\*\+\s\.,●■#]+|[•\-\*\+\s\.,●■#]+$', '', first_desc).strip()
                        desc_words = set(re.sub(r'[^a-zA-Z\s]', ' ', first_desc_clean).lower().split())
                        if desc_words & designation_keywords and len(first_desc_clean) < 60:
                            company = designation
                            designation = first_desc_clean
                            desc_lines = desc_lines[1:]

                # Apply semantic keyword checks on separate header lines
                # ONLY if _split_inline_job_header did NOT already set both designation and company
                if len(remaining_headers) >= 2 and (not designation or not company):
                    line1, line2 = remaining_headers[0], remaining_headers[1]
                    words1 = set(re.sub(r'[^a-zA-Z\s]', ' ', line1).lower().split())
                    words2 = set(re.sub(r'[^a-zA-Z\s]', ' ', line2).lower().split())
                    
                    is_l1_comp = bool(words1 & company_keywords)
                    is_l2_comp = bool(words2 & company_keywords)
                    is_l1_desig = bool(words1 & designation_keywords)
                    is_l2_desig = bool(words2 & designation_keywords)

                    if is_l1_desig and not is_l1_comp:
                        if not designation:
                            designation = line1
                        if not company:
                            company = line2
                    elif is_l2_desig and not is_l2_comp:
                        if not designation:
                            designation = line2
                        if not company:
                            company = line1
                    elif is_l1_comp and not is_l2_comp:
                        if not company:
                            company = line1
                        if not designation:
                            designation = line2
                    elif is_l2_comp and not is_l1_comp:
                        if not company:
                            company = line2
                        if not designation:
                            designation = line1

                if not company:
                    for header in remaining_headers:
                        words = set(re.sub(r'[^a-zA-Z\s]', ' ', header).lower().split())
                        if words & company_keywords and header != designation:
                            company = header
                            break
                if not company:
                    for header in remaining_headers:
                        if header != designation:
                            company = header
                            break
                if not designation:
                    designation = remaining_headers[0] if remaining_headers else ""

                # Check if any remaining header represents a city/location
                for header in remaining_headers:
                    if header != designation and header != company:
                        if any(loc.lower() in header.lower() for loc in locations) or "india" in header.lower():
                            location_val = header
                            break

                # Preserve description formatting exactly as parsed from the resume.
                # Bullet lines keep their original bullet character.
                # Paragraph lines stay as paragraphs. Never re-prefix every line with •.
                BULLET_CHARS = ('-', '\u2022', '*', '+', '\u25cf', '\u25a0', '\u25aa', '\u2013', '\u2192')
                desc_items = []
                for desc_line in desc_lines:
                    line_s = desc_line.strip()
                    # Skip nested section headings inside experience block
                    if line_s and any(h in line_s.lower() for h in ["achievements", "certifications", "competencies", "key skills"]):
                        continue
                    # Preserve original bullet/paragraph; only normalise corrupt multi-byte bullet artifacts
                    # These arise from Windows-1252/Latin-1 decoded as UTF-8 (â€¢ = •, etc.)
                    line_s = re.sub(r'â[^\s]{0,3}(?:\s|$)', '\u2022 ', line_s).strip()
                    desc_items.append(line_s)

                # Verify validity of job
                desig_words = set(re.sub(r'[^a-zA-Z\s]', ' ', designation).lower().split())
                comp_words = set(re.sub(r'[^a-zA-Z\s]', ' ', company).lower().split())
                if not start_date_val and not company:
                    continue
                if not start_date_val and not (desig_words & designation_keywords or comp_words & company_keywords):
                    continue

                experiences.append({
                    "designation": ResumeIntelligenceService._clean_job_header_value(designation),
                    "company": ResumeIntelligenceService._clean_job_header_value(company),
                    "location": ResumeIntelligenceService._clean_job_header_value(location_val),
                    "duration": ResumeIntelligenceService.get_duration_display(start_date_val, end_date_val) if start_date_val else "",
                    "description": "\n".join(desc_items),
                    "start_date": start_date_val,
                    "end_date": end_date_val
                })
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in work experience extraction: {str(e)}", exc_info=True)

        # 6. Parse Educations strictly from EDU section (Never classify institutes as degrees)
        educations = []
        try:
            edu_degree_keywords = {
                'mba', 'pgdm', 'b.com', 'bcom', 'b.tech', 'btech', 'b.e.', 'be', 'bsc', 'b.sc', 'msc', 'm.sc', 'phd', 
                'doctor', 'bachelor', 'master', 'diploma', 'high school', 'intermediate', '10th', '12th', 'ssc', 'hsc', 'school',
                'chartered accountant', 'accountant', 'ca', 'c.a.', 'bca', 'mca'
            }

            edu_blocks = []
            current_edu = []
            for line in sections["EDU"]:
                text_val = line["text"]
                if not text_val:
                    continue
                words = re.sub(r'[^a-zA-Z\s]', ' ', text_val).lower().split()
                has_deg = any(w in edu_degree_keywords for w in words)
                has_inst = any(w in ['university', 'college', 'school', 'institute', 'academy', 'icai', 'board'] for w in words)
                has_year = bool(re.search(r'\b(19\d\d|20\d\d)\b', text_val))

                is_new_edu = False
                if current_edu:
                    current_has_year = any(re.search(r'\b(19\d\d|20\d\d)\b', l["text"]) for l in current_edu)
                    current_has_deg = any(any(w in edu_degree_keywords for w in re.sub(r'[^a-zA-Z\s]', ' ', l["text"]).lower().split()) for l in current_edu)
                    if (has_deg and current_has_deg) or (has_year and current_has_year):
                        is_new_edu = True

                if is_new_edu:
                    edu_blocks.append(current_edu)
                    current_edu = [line]
                else:
                    current_edu.append(line)
            if current_edu:
                edu_blocks.append(current_edu)

            for block in edu_blocks:
                degree = ""
                institution = ""
                year = ""
                score = ""
                field = "General"
                
                block_texts = [b["text"] for b in block]
                full_text_edu = " ".join(block_texts)
                
                year_match = re.search(r'\b(19\d\d|20\d\d)\b', full_text_edu)
                if year_match:
                    year = year_match.group(0)
                    
                score_match = re.search(r'\b(\d{1,2}(?:\.\d{1,2})?\s*(?:%|cgpa|gpa|/10))\b', full_text_edu, re.I)
                if score_match:
                    score = score_match.group(1)
                else:
                    score_match2 = re.search(r'\b(gpa|cgpa|marks|percentage|score)[:\-\s]*(\d{1,2}(?:\.\d{1,2})?%?)\b', full_text_edu, re.I)
                    if score_match2:
                        score = score_match2.group(2)
                
                fields = ['computer science', 'information technology', 'software engineering', 'mechanical', 'electrical', 'electronics', 'civil', 'commerce', 'business', 'management', 'finance', 'marketing', 'accounting', 'arts', 'science', 'psychology']
                for f in fields:
                    if f in full_text_edu.lower():
                        field = f.title()
                        break

                degree_line = ""
                inst_line = ""

                for line in block:
                    line_clean = line["text"]
                    if not line_clean:
                        continue
                    words = re.sub(r'[^a-zA-Z\s]', ' ', line_clean).lower().split()
                    
                    # Check for institution keywords
                    is_inst = any(w in ['university', 'college', 'school', 'institute', 'academy', 'icai', 'board'] for w in words)
                    if is_inst and not inst_line:
                        inst_line = line_clean
                    
                    # Check for degree keywords (ignoring if it's already an institution line)
                    if not is_inst and not degree_line:
                        if any(w in edu_degree_keywords for w in words):
                            degree_line = line_clean

                if not degree_line and block:
                    # If the line is not an institution line, make it degree
                    other_lines = [l["text"] for l in block if l["text"] != inst_line]
                    if other_lines:
                        degree_line = other_lines[0]
                    else:
                        degree_line = block[0]["text"]
                
                if not inst_line and len(block) > 1:
                    other_lines = [l["text"] for l in block if l["text"] != degree_line]
                    if other_lines:
                        inst_line = other_lines[0]

                def clean_edu_field(val):
                    if not val:
                        return ""
                    val = val.replace('**', '').strip()
                    val = re.sub(r'\b(19\d\d|20\d\d)\b', '', val)
                    val = re.sub(r'\b(\d{1,2}(?:\.\d{1,2})?\s*(?:%|cgpa|gpa|/10))\b', '', val, flags=re.I)
                    val = re.sub(r'^[|/\-\s\.,:]+|[|/\-\s\.,:]+$', '', val).strip()
                    return " ".join(val.split())

                degree = clean_edu_field(degree_line) or "Degree"
                institution = clean_edu_field(inst_line) or "Institution"

                # Strict Degree vs Institution check: never classify ICAI/University as degree
                deg_lower = degree.lower()
                inst_lower = institution.lower()
                if any(w in deg_lower for w in ['university', 'college', 'school', 'institute', 'academy', 'icai']):
                    # Swap if degree has institute keywords and institution does not
                    if not any(w in inst_lower for w in ['university', 'college', 'school', 'institute', 'academy', 'icai']):
                        degree, institution = institution, degree
                
                inst_lower_check = institution.lower()
                if "icai" in inst_lower_check or "chartered accountant" in inst_lower_check or "institute of chartered" in inst_lower_check:
                    institution = "ICAI"
                
                if degree == "Institution":
                    degree = "Degree"

                educations.append({
                    "degree": degree,
                    "institution": institution,
                    "field_of_study": field,
                    "score": score or "N/A",
                    "start_date": f"{year}-01-01" if year.isdigit() else "2018-01-01",
                    "end_date": f"{year}-01-01" if year.isdigit() else "2022-01-01"
                })
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in education extraction: {str(e)}", exc_info=True)

        # 7. Parse Skills strictly from SKILLS section
        skills = []
        skill_groups = {}
        try:
            seen_skills = set()
            skill_groups = {
                "Technical Skills": [],
                "Programming Languages": [],
                "Frameworks": [],
                "Databases": [],
                "Cloud": [],
                "Tools": [],
                "Soft Skills": []
            }
            for line in sections["SKILLS"]:
                text_val = line["text"]
                if not text_val:
                    continue
                # Split on commas, pipes, semicolons, bullets, or dashes only when
                # surrounded by whitespace (to avoid splitting "Cross-Functional")
                parts = re.split(r'[,|;•●■]|\s+[–—]\s+', text_val)
                for part in parts:
                    s_clean = part.strip()
                    if not s_clean or len(s_clean) < 2 or len(s_clean) > 50:
                        continue
                    s_clean = re.sub(r'^[•\-\*\+\s\.,●■#]+|[•\-\*\+\s\.,●■#]+$', '', s_clean).strip()
                    if not s_clean:
                        continue
                    # Skip purely parenthetical fragments like "(MIS)" or single-word noise
                    if re.match(r'^\([^)]+\)$', s_clean):
                        # Merge with previous skill if available
                        if skills:
                            skills[-1] = f"{skills[-1]} {s_clean}"
                        continue
                    s_lower = s_clean.lower()
                    if s_lower not in seen_skills:
                        seen_skills.add(s_lower)
                        skills.append(s_clean)
                        group_name = ResumeIntelligenceService._skill_group_for(s_clean)
                        skill_groups[group_name].append(s_clean)

            skill_groups = {group: values for group, values in skill_groups.items() if values}
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in skills extraction: {str(e)}", exc_info=True)

        # 8. Parse Certifications strictly from CERT section
        certs = []
        try:
            cert_blocks = []
            current_cert = []
            for line in sections["CERT"]:
                text_val = line["text"]
                if not text_val:
                    continue
                if len(current_cert) >= 2:
                    cert_blocks.append(current_cert)
                    current_cert = []
                current_cert.append(text_val)
            if current_cert:
                cert_blocks.append(current_cert)
                
            for block in cert_blocks:
                c_name = block[0] if block else "Certification"
                org = block[1] if len(block) > 1 else "Accredited Body"
                certs.append({
                    "name": c_name[:100],
                    "issuing_organization": org,
                    "issue_date": "2023-06-01"
                })
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in certifications extraction: {str(e)}", exc_info=True)

        # 9. Parse Projects strictly from PROJECT section
        projs = []
        try:
            proj_blocks = []
            current_proj = []
            for line in sections["PROJECT"]:
                text_val = line["text"]
                if not text_val:
                    if current_proj:
                        current_proj.append(text_val)
                    continue
                if text_val.startswith(('-', '•', '*', '+', '●', '■')) and current_proj:
                    current_proj.append(text_val)
                else:
                    if len(current_proj) >= 2:
                        proj_blocks.append(current_proj)
                        current_proj = []
                    current_proj.append(text_val)
            if current_proj:
                proj_blocks.append(current_proj)
                
            for block in proj_blocks:
                title = block[0] if block else "Project"
                desc = "\n".join(block[1:]) if len(block) > 1 else title
                projs.append({
                    "title": title[:100],
                    "description": desc,
                    "link": ""
                })
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in projects extraction: {str(e)}", exc_info=True)

        # 10. Parse Achievements strictly from ACHIEVEMENTS section
        achievements = []
        try:
            ach_text_list = []
            for line in sections.get("ACHIEVEMENTS", []):
                ach_text_list.append(line["text"])
            if ach_text_list:
                ach_html = ResumeIntelligenceService.parse_experience_description_to_html("\n".join(ach_text_list))
                achievements = [ach_html]
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in achievements extraction: {str(e)}", exc_info=True)

        # Calculate experience years strictly from parsed start/end dates
        total_exp = 0.0
        try:
            for exp in experiences:
                s_date = exp.get("start_date")
                e_date = exp.get("end_date")
                if s_date:
                    total_exp += ResumeIntelligenceService.calculate_experience_years_from_dates(s_date, e_date)
            total_exp = round(total_exp, 1)

            # Fallback/refinement using text-based experience extraction
            text_exp = 0.0
            m = re.search(r'\b(\d+)\s*years?\s*(?:\d+\s*months?)?\b', text, re.I)
            if m:
                val = float(m.group(1))
                if 0.5 <= val <= 40.0:
                    text_exp = val
            else:
                m = re.search(r'\b(\d+)\s*years?\s+(?:of\s+)?(?:hands-on\s+)?experience\b', text, re.I)
                if m:
                    val = float(m.group(1))
                    if 0.5 <= val <= 40.0:
                        text_exp = val

            if text_exp > 0.0:
                if total_exp == 0.0 or abs(total_exp - text_exp) <= 2.0:
                    total_exp = text_exp
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in experience calculation: {str(e)}", exc_info=True)

        try:
            # Print raw extracted experience JSON and final saved experience JSON (for user logs)
            print("--- [RAW EXTRACTED EXPERIENCE JSON] ---")
            print(json.dumps(experiences, indent=2))
            print("--- [FINAL SAVED EXPERIENCE JSON] ---")
            print(json.dumps(experiences, indent=2))
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Error in experience JSON logging: {str(e)}", exc_info=True)

        return {
            "personal_info": {
                "name": name or "John Doe",
                "email": email or "candidate@example.com",
                "phone": phone or "9876543210",
                "location": location,
                "address": address,
                "city": city,
                "linkedin_url": linkedin,
                "portfolio_url": portfolio,
                "current_company": experiences[0]["company"] if experiences else "",
                "current_designation": experiences[0]["designation"] if experiences else "",
                "total_experience": total_exp
            },
            "summary": summary,
            "skills": skills,
            "skill_groups": skill_groups,
            "education": educations,
            "experience": experiences,
            "projects": projs,
            "certifications": certs,
            "achievements": achievements,
            "languages": ["English"],
            "metadata": {
                "parsed_at": datetime.now().isoformat(),
                "word_count": len(text.split())
            }
        }

    @staticmethod
    def ai_improve_resume_data(data: dict) -> dict:
        """
        AI Assist processor that cleans common OCR typos, deduplicates skills,
        normalizes entities, suggests missing items, and rewrites experiences
        using professional ATS STAR format.
        """
        improved = json.loads(json.dumps(data)) # Deep copy

        # 1. Clean Name/Contact
        info = improved["personal_info"]
        info["name"] = info["name"].strip().title()
        
        # 2. Normalize company names
        for exp in improved.get("experience", []):
            comp = exp.get("company", "")
            exp["company"] = comp.title()
            
            # Preserve the exact description formatting produced by parse_resume_nlp.
            # Do NOT strip bullets and re-prefix them — that destroys paragraph formatting.
            desc = exp.get("description", "")
            # Do not filter out blank lines to preserve original spacing
            cleaned_desc_lines = [l for l in desc.split('\n')]
            exp["description"] = "\n".join(cleaned_desc_lines)

        # Update current designation and company in info dictionary if experiences exist
        if improved.get("experience"):
            first_exp = improved["experience"][0]
            info["current_company"] = first_exp.get("company", "")
            info["current_designation"] = first_exp.get("designation", "")

        # Improve current designation using summary context if available
        current_desig = info.get('current_designation') or ''
        summary_clean = data.get("summary", "")
        if summary_clean and current_desig:
            summary_clean = " ".join(summary_clean.split())
            # Look for patterns like "experience as a Marketing Executive and in Floor Management"
            match = re.search(r'experience as a?\s*([^,.]+?)\s+and\s+(?:in\s+)?([^,.]+)', summary_clean, re.I)
            if match:
                role1 = match.group(1).strip()
                role2 = match.group(2).strip()
                if len(role1) < 50 and len(role2) < 50:
                    if role1.lower() in current_desig.lower() or current_desig.lower() in role1.lower():
                        combined = f"{role1} / {role2}"
                        combined = " / ".join([r.strip().title() for r in combined.split('/')])
                        info["current_designation"] = combined
                        if improved.get("experience"):
                            improved["experience"][0]["designation"] = combined

        # 3. Normalize degrees
        for edu in improved.get("education", []):
            deg = edu.get("degree", "").lower()
            if "btech" in deg or "b.tech" in deg or "b.e." in deg or "bachelor of technology" in deg:
                edu["degree"] = "B.Tech"
            elif "mtech" in deg or "m.tech" in deg or "master" in deg:
                edu["degree"] = "M.Tech"
            elif "phd" in deg or "doctor" in deg:
                edu["degree"] = "Ph.D."
            elif "diploma" in deg:
                edu["degree"] = "Diploma"
            elif "intermediate" in deg or "12th" in deg or "hsc" in deg:
                edu["degree"] = "Intermediate"
            elif "high school" in deg or "10th" in deg or "ssc" in deg or "school" in deg:
                edu["degree"] = "High School"
            inst = edu.get("institution", "").strip()
            if "icai" in inst.lower():
                edu["institution"] = "ICAI"
            else:
                edu["institution"] = inst.title()

        # 4. Generate professional Summary if not present or improve it
        current_desig = info.get('current_designation') or 'Professional'
        if current_desig.lower() == 'professional':
            current_desig = 'Professional'
            
        orig_summary = data.get("summary", "")
        cand_name = info.get("name", "")
        
        if orig_summary:
            clean_summary = orig_summary
            if cand_name:
                clean_summary = re.sub(rf'\b{re.escape(cand_name)}\b', '', clean_summary, flags=re.I).strip()
                clean_summary = re.sub(r'^\s*[-–—,.:;]\s*', '', clean_summary)
            improved["summary"] = clean_summary
        else:
            skills_str = ", ".join(improved.get("skills", [])[:5])
            improved["summary"] = f"Results-driven {current_desig} with {info.get('total_experience', 2)} years of experience. Highly skilled in {skills_str}, with a proven track record of delivering high-quality results."

        # 5. Deduplicate skills (Do not suggest fake/missing skills)
        skills_set = {s.strip().title() for s in improved.get("skills", [])}
        improved["skills"] = sorted(list(skills_set))

        # 6. Certifications (Do not suggest fake/missing certifications)
        improved["certifications"] = improved.get("certifications", [])

        return improved

    @staticmethod
    def calculate_duplicate_similarity(c1: CandidateProfile, c2: CandidateProfile) -> dict:
        """
        Candidate duplicate detection using Name similarity, Jaccard token cosine overlap,
        Email, Phone, and LinkedIn match validation.
        """
        reasons = []
        is_duplicate = False
        score = 0

        # Exact identifiers match
        if c1.user.email and c2.user.email and c1.user.email.lower() == c2.user.email.lower():
            reasons.append("Exact Email Match")
            score = 100
            is_duplicate = True

        if c1.user.phone_number and c2.user.phone_number and c1.user.phone_number == c2.user.phone_number:
            reasons.append("Exact Phone Match")
            score = 98
            is_duplicate = True

        if c1.linkedin_url and c2.linkedin_url and c1.linkedin_url.strip().lower() == c2.linkedin_url.strip().lower():
            reasons.append("Exact LinkedIn Match")
            score = 95
            is_duplicate = True

        if c1.sha256 and c2.sha256 and c1.sha256 == c2.sha256:
            reasons.append("Exact Resume Hash Match")
            score = 100
            is_duplicate = True

        # Name similarity (Jaccard similarity of name words)
        n1 = set((c1.full_name or "").lower().split())
        n2 = set((c2.full_name or "").lower().split())
        name_sim = 0
        if n1 and n2:
            name_sim = len(n1.intersection(n2)) / len(n1.union(n2))
            if name_sim >= 0.75:
                reasons.append(f"Name Match Similarity ({int(name_sim * 100)}%)")
                score = max(score, int(name_sim * 90))

        # Profile content cosine word similarity
        text1 = f"{c1.summary} {' '.join(s.skill_name for s in c1.skills.all())}".lower()
        text2 = f"{c2.summary} {' '.join(s.skill_name for s in c2.skills.all())}".lower()
        
        words1 = re.findall(r'\w+', text1)
        words2 = re.findall(r'\w+', text2)
        
        from collections import Counter
        c_words1 = Counter(words1)
        c_words2 = Counter(words2)
        
        # Calculate cosine similarity
        intersection = set(c_words1.keys()) & set(c_words2.keys())
        numerator = sum([c_words1[x] * c_words2[x] for x in intersection])
        
        sum1 = sum([c_words1[x]**2 for x in c_words1.keys()])
        sum2 = sum([c_words2[x]**2 for x in c_words2.keys()])
        denominator = math.sqrt(sum1) * math.sqrt(sum2)
        
        cosine_sim = numerator / denominator if denominator else 0.0
        if cosine_sim >= 0.70:
            reasons.append(f"Profile Text Cosine Similarity ({int(cosine_sim * 100)}%)")
            score = max(score, int(cosine_sim * 100))

        return {
            "is_duplicate": is_duplicate,
            "similarity_score": score,
            "reasons": reasons,
            "duplicate_candidate_id": str(c2.id),
            "duplicate_candidate_name": c2.full_name or c2.user.email
        }

    @staticmethod
    def generate_ats_friendly_pdf(candidate: CandidateProfile) -> bytes:
        """
        Compiles candidate database records into a clean, ATS-compliant, print-friendly
        PDF document using ReportLab Flowables.
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()
        
        # Custom stylesheet to avoid collision
        title_style = ParagraphStyle(
            'ResumeTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=22,
            leading=26,
            textColor=colors.HexColor('#111827'),
            spaceAfter=6
        )
        
        subtitle_style = ParagraphStyle(
            'ResumeSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#4B5563'),
            spaceAfter=15
        )

        section_heading = ParagraphStyle(
            'ResumeSectionHeading',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=13,
            leading=16,
            textColor=colors.HexColor('#1E3A8A'),
            spaceBefore=12,
            spaceAfter=6
        )

        body_style = ParagraphStyle(
            'ResumeBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=14,
            textColor=colors.HexColor('#374151'),
            spaceAfter=4
        )

        bullet_style = ParagraphStyle(
            'ResumeBullet',
            parent=body_style,
            leftIndent=15,
            bulletIndent=5,
            spaceAfter=3
        )

        story = []

        # 1. Header Information
        full_name = escape_plain_text(candidate.full_name or candidate.user.email)
        story.append(Paragraph(full_name, title_style))
        contact_info = []
        if candidate.user.email:
            contact_info.append(candidate.user.email)
        if candidate.user.phone_number:
            contact_info.append(candidate.user.phone_number)
        if candidate.location:
            contact_info.append(candidate.location)
        if candidate.linkedin_url:
            contact_info.append("LinkedIn")
            
        escaped_contact_info = [escape_plain_text(info) for info in contact_info]
        story.append(Paragraph(" | ".join(escaped_contact_info), subtitle_style))

        # 2. Professional Summary
        if candidate.summary:
            story.append(Paragraph("PROFESSIONAL SUMMARY", section_heading))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E5E7EB'), spaceBefore=2, spaceAfter=8))
            story.append(Paragraph(sanitize_html_for_reportlab(candidate.summary), body_style))
            story.append(Spacer(1, 10))

        # 3. Core Skills
        skills = candidate.skills.all()
        if skills.exists():
            story.append(Paragraph("CORE SKILLS", section_heading))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E5E7EB'), spaceBefore=2, spaceAfter=8))
            skills_text = ", ".join([escape_plain_text(s.skill_name) for s in skills])
            story.append(Paragraph(skills_text, body_style))
            story.append(Spacer(1, 10))

        # 4. Work Experience
        exps = candidate.experiences.all()
        if exps.exists():
            story.append(Paragraph("WORK EXPERIENCE", section_heading))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E5E7EB'), spaceBefore=2, spaceAfter=8))
            for exp in exps:
                dates = f"{exp.start_date.strftime('%b %Y') if exp.start_date else ''} - {exp.end_date.strftime('%b %Y') if exp.end_date else 'Present'}"
                heading = f"<b>{escape_plain_text(exp.designation)}</b> @ {escape_plain_text(exp.company_name)} ({dates})"
                story.append(Paragraph(heading, body_style))
                
                # Split description into bullets if possible
                bullets = extract_bullets_from_html(exp.description)
                for bullet in bullets:
                    b_text = bullet.strip().lstrip('-•*').strip()
                    if b_text:
                        story.append(Paragraph(f"• {sanitize_html_for_reportlab(b_text)}", bullet_style))
                story.append(Spacer(1, 6))
            story.append(Spacer(1, 10))

        # 5. Education
        edus = candidate.educations.all()
        if edus.exists():
            story.append(Paragraph("EDUCATION", section_heading))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E5E7EB'), spaceBefore=2, spaceAfter=8))
            for edu in edus:
                dates = f"{edu.start_date.strftime('%Y') if edu.start_date else ''} - {edu.end_date.strftime('%Y') if edu.end_date else ''}"
                field = f" in {escape_plain_text(edu.field_of_study)}" if edu.field_of_study else ""
                edu_text = f"<b>{escape_plain_text(edu.degree)}{field}</b> - {escape_plain_text(edu.institution)} ({dates})"
                story.append(Paragraph(edu_text, body_style))
            story.append(Spacer(1, 10))

        # 6. Projects
        projs = candidate.projects.all()
        if projs.exists():
            story.append(Paragraph("PROJECTS", section_heading))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E5E7EB'), spaceBefore=2, spaceAfter=8))
            for proj in projs:
                p_text = f"<b>{escape_plain_text(proj.title)}</b>"
                if proj.link:
                    escaped_link = escape_plain_text(proj.link)
                    p_text += f" (<a href='{escaped_link}'>{escaped_link}</a>)"
                story.append(Paragraph(p_text, body_style))
                story.append(Paragraph(sanitize_html_for_reportlab(proj.description), bullet_style))
                story.append(Spacer(1, 4))
            story.append(Spacer(1, 10))

        # 7. Certifications
        certs = candidate.certifications.all()
        if certs.exists():
            story.append(Paragraph("CERTIFICATIONS", section_heading))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E5E7EB'), spaceBefore=2, spaceAfter=8))
            for cert in certs:
                org = f" - {escape_plain_text(cert.issuing_organization)}" if cert.issuing_organization else ""
                date = f" ({cert.issue_date.strftime('%B %Y')})" if cert.issue_date else ""
                story.append(Paragraph(f"• {escape_plain_text(cert.name)}{org}{date}", bullet_style))

        # Render PDF document
        doc.build(story)
        pdf_data = buffer.getvalue()
        buffer.close()
        return pdf_data
